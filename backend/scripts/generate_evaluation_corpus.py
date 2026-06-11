"""Generate a reproducible 30-document stratified synthetic evaluation corpus."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from generate_sample_docx import add_native_formula, add_picture, configure_document


ROOT = Path(__file__).resolve().parents[2]
SAMPLE_DIR = ROOT / "sample_documents" / "evaluation"
GOLDEN_DIR = ROOT / "backend" / "evaluation" / "goldens"
MANIFEST_PATH = ROOT / "backend" / "evaluation" / "manifest.json"
CATEGORIES = (
    ("zh", "中文普通论文"),
    ("en", "英文普通论文"),
    ("figure", "图表密集论文"),
    ("formula", "公式密集论文"),
    ("reference", "参考文献复杂论文"),
    ("anomaly", "异常排版论文"),
)


def add_title(document: Document, text: str, *, obvious: bool = True) -> None:
    paragraph = document.add_paragraph()
    if obvious:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
    else:
        paragraph.add_run(text)


def corrected_article(golden: dict, index: int, language: str) -> dict:
    return {
        "title": golden["title"],
        "doi": f"10.9999/word2jats.eval.{index:02d}",
        "article_type": "research-article",
        "lang": language,
        "journal_title": "Word2JATS Evaluation Journal",
        "journal_id": "W2J-EVAL",
        "issn": "2099-9999",
        "publisher_name": "Word2JATS Publishing Lab",
        "subject": "Structured publishing",
        "pub_year": "2026",
        "pub_month": "06",
        "pub_day": "11",
        "authors": [{
            "name": "Alice Smith" if language == "en" else "张三",
            "orcid": f"0000-0002-0000-{index:04d}",
            "affiliation_ids": ["aff1"],
        }],
        "affiliations": ["Word2JATS Publishing Lab"],
        "abstract": golden["abstract"],
        "keywords": golden["keywords"],
        "sections": [
            {
                "title": section["title"],
                "level": section.get("level", 1),
                "paragraphs": section.get("paragraphs", ["人工校正后的正文内容。"]),
            }
            for section in golden["sections"]
        ],
        "figures": [
            {"id": f"fig{i}", "caption": f"Figure {i}", "path": "", "section_index": 0}
            for i, _ in enumerate(golden.get("figures", []), start=1)
        ],
        "tables": [
            {
                "id": f"tab{i}", "caption": f"Table {i}",
                "rows": [["Metric", "Value"], ["Accuracy", "95%"]], "section_index": 0,
            }
            for i, _ in enumerate(golden.get("tables", []), start=1)
        ],
        "lists": [],
        "formulas": [
            {"id": f"eq{i}", "content": f"x_{i} = y_{i}", "latex": f"x_{i} = y_{i}", "section_index": 0}
            for i, _ in enumerate(golden.get("formulas", []), start=1)
        ],
        "references": [
            {"id": f"ref{i}", "label": f"[{i}]", "raw": f"Evaluation reference {i}. 2026."}
            for i, _ in enumerate(golden.get("references", []), start=1)
        ],
    }


def base_document(title: str, abstract: str, keywords: list[str], *, english: bool = False, obvious_title: bool = True) -> Document:
    document = Document()
    configure_document(document)
    add_title(document, title, obvious=obvious_title)
    document.add_paragraph("Alice Smith, Bob Lee" if english else "张三，李四")
    document.add_paragraph("School of Publishing, Demo University" if english else "示范大学 数字出版学院")
    document.add_paragraph(("Abstract: " if english else "摘要：") + abstract)
    document.add_paragraph(("Keywords: " if english else "关键词：") + "; ".join(keywords))
    return document


def build_sample(kind: str, index: int) -> tuple[Document, dict, str]:
    english = kind == "en"
    language = "en" if english else "zh"
    title = (
        f"Structured Publishing Evaluation Study {index}"
        if english else f"结构化出版评测研究 {index}"
    )
    abstract = (
        f"This evaluation manuscript examines reproducible JATS conversion scenario {index}."
        if english else f"本文用于评估第 {index} 种可复现 JATS 结构化转换场景。"
    )
    keywords = ["JATS", "Word", f"Case{index}"] if english else ["JATS", "Word", f"场景{index}"]
    document = base_document(title, abstract, keywords, english=english)
    sections = []
    figures: list[dict] = []
    tables: list[dict] = []
    formulas: list[dict] = []
    references: list[dict] = []

    if kind in {"zh", "en"}:
        headings = ["Introduction", "Method", "Conclusion"] if english else ["引言", "方法", "结论"]
        for number, heading in enumerate(headings, start=1):
            document.add_paragraph(f"{number} {heading}")
            document.add_paragraph(f"Section {number} content for evaluation." if english else f"第 {number} 节评测正文内容。")
            sections.append({"title": heading, "paragraphs": ["正文"]})
    elif kind == "figure":
        document.add_paragraph("1 图表实验")
        document.add_paragraph("本节包含多张图片与多个表格。")
        sections.append({"title": "图表实验", "paragraphs": ["正文"]})
        for number in range(1, index + 2):
            add_picture(document, f"eval_fig_{index}_{number}.png", (20 * number, 90, 130), number % 4)
            if not (index == 5 and number == index + 1):
                document.add_paragraph(f"图{number} 图表密集样本 {number}")
            figures.append({})
        for number in range(1, 3):
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text, table.rows[0].cells[1].text = "指标", "值"
            table.rows[1].cells[0].text, table.rows[1].cells[1].text = "准确率", f"{90 + number}%"
            document.add_paragraph(f"表{number} 实验结果")
            tables.append({})
    elif kind == "formula":
        document.add_paragraph("1 公式方法")
        document.add_paragraph("下列段落用于评估公式识别。")
        sections.append({"title": "公式方法", "paragraphs": ["正文"]})
        formula_texts = ["x = α + β", "∑ x_i ≤ λ", r"\frac{a}{b} = \sqrt{c}", "lim x→0 sin(x)/x = 1"]
        for text in formula_texts[: min(4, index)]:
            document.add_paragraph(text)
            formulas.append({})
        if index >= 3:
            add_native_formula(document)
            formulas.append({})
        if index == 5:
            document.add_paragraph("矩阵 A 的特征值需要人工复核")
            formulas.append({})
    elif kind == "reference":
        document.add_paragraph("1 文献综述")
        document.add_paragraph("正文引用参考文献 [1] 和 [2]。")
        sections.append({"title": "文献综述", "paragraphs": ["正文"]})
    else:
        document = Document()
        configure_document(document)
        document.add_paragraph(f"内部编号 EVAL-{index}")
        add_title(document, title, obvious=index % 2 == 0)
        document.add_paragraph("张三 李四")
        document.add_paragraph("示范大学")
        marker = "内容摘要：" if index in {1, 3, 5} else "摘要："
        document.add_paragraph(marker + abstract)
        keyword_marker = "主题词：" if index in {2, 4} else "关键词："
        document.add_paragraph(keyword_marker + "；".join(keywords))
        document.add_paragraph("章节一 方法" if index == 5 else "1 方法")
        document.add_paragraph("异常排版条件下的正文。")
        sections.append({"title": "方法", "paragraphs": ["正文"]})

    if kind not in {"reference"}:
        document.add_paragraph("参考文献" if not english else "References")
    else:
        document.add_paragraph("参考文献")
    reference_lines = [
        "[1] Zhang S, Li Q. Structured publishing[J]. Journal of XML, 2025, 12(3): 10-18. doi:10.1234/eval.1",
        "2. 王五. 数字出版结构化方法[J]. 出版科学, 2026, 34(1): 20-28.",
        "（3）National Information Standards Organization. JATS: Journal Article Tag Suite[S].",
    ]
    count = 3 if kind == "reference" else 1
    if kind == "reference" and index >= 4:
        reference_lines.append("Unnumbered and irregular citation without a standard year")
        count = 4
    for line in reference_lines[:count]:
        document.add_paragraph(line)
        references.append({})

    golden = {
        "category": dict(CATEGORIES)[kind],
        "title": title,
        "abstract": abstract,
        "keywords": keywords,
        "sections": sections,
        "figures": figures,
        "tables": tables,
        "formulas": formulas,
        "references": references,
    }
    golden["corrected_article"] = corrected_article(golden, index + list(dict(CATEGORIES)).index(kind) * 5, language)
    return document, golden, language


def generate() -> list[dict]:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    GOLDEN_DIR.mkdir(parents=True, exist_ok=True)
    for path in GOLDEN_DIR.glob("eval_*.json"):
        path.unlink()
    for path in SAMPLE_DIR.glob("eval_*.docx"):
        path.unlink()

    manifest = []
    for kind, category in CATEGORIES:
        for index in range(1, 6):
            stem = f"eval_{kind}_{index:02d}"
            document, golden, language = build_sample(kind, index)
            document.save(SAMPLE_DIR / f"{stem}.docx")
            (GOLDEN_DIR / f"{stem}.json").write_text(
                json.dumps(golden, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            manifest.append({
                "sample": f"evaluation/{stem}.docx",
                "golden": f"{stem}.json",
                "category": category,
                "language": language,
                "difficulty": "hard" if kind in {"anomaly", "reference"} and index >= 4 else "standard",
            })
    MANIFEST_PATH.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


if __name__ == "__main__":
    entries = generate()
    print(f"Generated {len(entries)} evaluation documents.")
