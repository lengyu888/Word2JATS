"""Generate the single, comprehensive Word2JATS acceptance document."""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "sample_documents"
FINAL_OUTPUT = OUTPUT_DIR / "word2jats_final_acceptance.docx"
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def make_png(path: Path, color: tuple[int, int, int], label_band: int) -> None:
    width, height = 520, 190
    rows = []
    for y in range(height):
        row = bytearray([0])
        for x in range(width):
            in_card = 30 < x < 490 and 30 < y < 160
            in_band = 55 + label_band * 20 < x < 215 + label_band * 20 and 70 < y < 120
            pixel = color if in_card else (243, 240, 230)
            if in_band:
                pixel = (255, 255, 255)
            row.extend((*pixel, 255))
        rows.append(bytes(row))

    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
        )

    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(b"".join(rows)))
        + chunk(b"IEND", b"")
    )


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(document: Document, text: str, *, obvious: bool = True) -> None:
    paragraph = document.add_paragraph()
    if obvious:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = obvious
    run.font.size = Pt(18 if obvious else 12)


def add_picture(
    document: Document, filename: str, color: tuple[int, int, int], band: int
) -> None:
    image_path = OUTPUT_DIR / filename
    make_png(image_path, color, band)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.3))
    image_path.unlink()


def run(text: str) -> str:
    return f"<m:r><m:t>{text}</m:t></m:r>"


def add_native_formula(document: Document, content: str) -> None:
    paragraph = document.add_paragraph()
    paragraph._p.append(parse_xml(f'<m:oMath xmlns:m="{OMML_NS}">{content}</m:oMath>'))


def add_basic_formula(document: Document) -> None:
    add_native_formula(
        document,
        f"""
        <m:f><m:num>{run("a")}</m:num><m:den>{run("b")}</m:den></m:f>
        {run("+")}
        <m:sSup><m:e>{run("x")}</m:e><m:sup>{run("2")}</m:sup></m:sSup>
        {run("+")}
        <m:rad><m:e>{run("z")}</m:e></m:rad>
        """,
    )


def add_matrix_formula(document: Document) -> None:
    add_native_formula(
        document,
        f"""
        <m:m>
          <m:mr><m:e>{run("a")}</m:e><m:e>{run("b")}</m:e></m:mr>
          <m:mr><m:e>{run("c")}</m:e><m:e>{run("d")}</m:e></m:mr>
        </m:m>
        """,
    )


def add_cases_formula(document: Document) -> None:
    add_native_formula(
        document,
        f"""
        <m:d>
          <m:dPr><m:begChr m:val="{{"/><m:endChr m:val=""/></m:dPr>
          <m:e><m:eqArr>
            <m:e>{run("x, x&gt;0")}</m:e>
            <m:e>{run("-x, x≤0")}</m:e>
          </m:eqArr></m:e>
        </m:d>
        """,
    )


def add_nary_and_accents(document: Document) -> None:
    add_native_formula(
        document,
        f"""
        <m:nary>
          <m:naryPr><m:chr m:val="∑"/></m:naryPr>
          <m:sub>{run("i=1")}</m:sub><m:sup>{run("n")}</m:sup><m:e>{run("x_i")}</m:e>
        </m:nary>
        {run("+")}
        <m:nary>
          <m:naryPr><m:chr m:val="∫"/></m:naryPr>
          <m:sub>{run("0")}</m:sub><m:sup>{run("1")}</m:sup><m:e>{run("f(x)dx")}</m:e>
        </m:nary>
        {run("+")}
        <m:acc><m:accPr><m:chr m:val="^"/></m:accPr><m:e>{run("x")}</m:e></m:acc>
        """,
    )


def add_partial_formula(document: Document) -> None:
    add_native_formula(
        document,
        f"""
        <m:acc><m:accPr><m:chr m:val="⃛"/></m:accPr><m:e>{run("q")}</m:e></m:acc>
        <m:unknownComplex>{run("kept")}</m:unknownComplex>
        """,
    )


def build_final_document() -> Document:
    document = Document()
    configure_document(document)
    add_title(document, "Word2JATS：学术期刊智能结构化转换全流程验收稿")
    document.add_paragraph("张三，李四，王五")
    document.add_paragraph("未来出版大学 数字出版学院，智能内容处理实验室")
    document.add_paragraph(
        "摘要：本文构造一篇覆盖 Word2JATS 决赛展示能力的学术论文验收稿，"
        "用于验证 DOCX 真实文档流解析、结构化 JSON、JATS XML、正式 Schema 校验和质量报告。"
    )
    document.add_paragraph(
        "文档还包含图片、表格、列表、正文交叉引用、Word 原生公式和细粒度参考文献，"
        "可用于演示人工校正后重新生成 XML 与 ZIP 交付闭环。"
    )
    document.add_paragraph("关键词：JATS；Word；结构化出版；OMML；质量校验")

    document.add_paragraph("1 引言")
    document.add_paragraph(
        "学术出版需要将作者提交的 Word 稿件转换为机器可读、可交换和可长期保存的结构化内容。"
    )
    document.add_paragraph(
        "如图1和表1所示，系统按真实文档顺序恢复出版对象；公式结果见式（1），"
        "相关标准与研究见参考文献[1,2]。"
    )

    document.add_paragraph("1.1 系统能力")
    document.add_paragraph("（1）解析标题、作者、单位、摘要、关键词和多级章节")
    document.add_paragraph("（2）绑定图片、图题、表格、表题和所属章节")
    document.add_paragraph("（3）转换 OMML、恢复交叉引用并执行 JATS Schema 校验")

    document.add_paragraph("2 文档流与图表解析")
    add_picture(document, "final_architecture.png", (0, 109, 119), 0)
    document.add_paragraph("图1 Word2JATS 系统总体架构")
    add_picture(document, "final_workflow.png", (199, 138, 30), 1)
    document.add_paragraph("Figure 2 Conversion workflow")
    document.add_paragraph("正文再次引用图 1和Figure 2，用于验证图引用反向统计。")

    document.add_paragraph("表1 核心模块与输出")
    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    values = [
        ["模块", "输入", "输出"],
        ["文档流解析", "DOCX", "结构化 JSON"],
        ["JATS 生成", "Article JSON", "article.xml"],
        ["质量校验", "XML", "质量报告"],
    ]
    for row, values_row in zip(table.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value

    document.add_paragraph("Table 2 Quality score example")
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    values = [["Metric", "Value"], ["XML well-formed", "100%"], ["Schema", "review"]]
    for row, values_row in zip(table.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = value
    document.add_paragraph("正文中的表 1与Table 2均可在图表预览页查看对应 JATS 片段。")

    document.add_paragraph("3 Word 原生公式转换")
    document.add_paragraph("式（1）包含分数、上标和根号：")
    add_basic_formula(document)
    document.add_paragraph("式（2）为 2×2 矩阵：")
    add_matrix_formula(document)
    document.add_paragraph("式（3）为分段函数：")
    add_cases_formula(document)
    document.add_paragraph("式（4）包含带上下限的大运算符和常见重音：")
    add_nary_and_accents(document)
    document.add_paragraph("式（5）故意包含未完整支持的复杂重音，用于验证 partial 稳定降级：")
    add_partial_formula(document)

    document.add_paragraph("4 质量闭环与交付")
    document.add_paragraph(
        "系统输出结构化 JSON、JATS XML、校验结果、质量分、原文映射视图和图表预览。"
    )
    document.add_paragraph(
        "人工补充 ORCID、ISSN、DOI 等真实出版元数据后，可重新生成 XML 并再次执行 Schema 校验。"
    )
    document.add_paragraph(
        "单篇结果可导出为包含 article.xml、article.json、校验报告、质量报告和媒体文件的 ZIP 包。"
    )

    document.add_paragraph("5 结论")
    document.add_paragraph(
        "本验收稿作为仓库中唯一的人工测试 Word 文档，覆盖系统主要成功路径与可控降级路径。"
    )

    document.add_paragraph("参考文献")
    document.add_paragraph(
        "[1] 张三, 李四. 学术出版结构化转换技术研究[J]. 数字出版, 2025, 10(2): 1-8. "
        "doi:10.1234/w2j.2025.001"
    )
    document.add_paragraph(
        "2. National Information Standards Organization. JATS: Journal Article Tag Suite[S]. 2024."
    )
    document.add_paragraph(
        "（3）王五. Word 文档智能解析方法研究[J]. 出版科学, 2026, 34(1): 20-28."
    )
    return document


def generate(output: Path = FINAL_OUTPUT) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    build_final_document().save(output)
    return output


if __name__ == "__main__":
    print(generate())
