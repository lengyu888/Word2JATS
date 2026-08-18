from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from lxml import etree

from app.services.docx_parser import DocxParser
from app.services.jats_generator import JatsGenerator
from app.services.validator import ArticleValidator
from app.services.document_flow_parser import DocumentFlowParser
from tests.test_services import make_png


def append_omml(paragraph, text: str) -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    math.append(run)
    paragraph._p.append(math)


def test_multiline_title_precedes_initialed_author_list(tmp_path):
    path = tmp_path / "multiline-title.docx"
    document = Document()
    title = document.add_paragraph("Investigating Solvent Extraction")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Using a Sustainable Extractant")
    document.add_paragraph("F. Smith1*, K. Jones2")
    document.add_paragraph("1 Department of Chemistry, Example University")
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: solvent; extraction; JATS")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["title"] == (
        "Investigating Solvent Extraction Using a Sustainable Extractant"
    )
    assert [author["name"] for author in article["authors"]] == [
        "F. Smith", "K. Jones"
    ]


def test_legacy_heading_style_offset_and_long_body_recovery(tmp_path):
    path = tmp_path / "legacy-heading-offset.docx"
    document = Document()
    title = document.add_paragraph("Legacy Heading Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: legacy; headings; JATS")
    level_two = document.styles.add_style("2", 1)
    level_three = document.styles.add_style("3", 1)
    level_four = document.styles.add_style("4", 1)
    heading = document.add_paragraph("Introduction")
    heading.style = level_two
    subheading = document.add_paragraph("Experimental Design")
    subheading.style = level_three
    long_body = document.add_paragraph(" ".join(["Detailed body evidence"] * 35))
    long_body.style = level_four
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [(section["title"], section["level"]) for section in article["sections"]] == [
        ("Introduction", 1), ("Experimental Design", 2)
    ]
    assert article["sections"][1]["paragraphs"] == [long_body.text]


def test_named_heading_styles_restore_nested_section_levels(tmp_path):
    path = tmp_path / "named-heading-levels.docx"
    document = Document()
    title = document.add_paragraph("Named Heading Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: headings; hierarchy; JATS")
    document.add_heading("Introduction", level=1)
    document.add_heading("Experimental Design", level=2)
    document.add_heading("Sampling", level=3)
    document.add_paragraph("Body evidence.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["level"] for section in article["sections"]] == [1, 2, 3]
    assert [section["label"] for section in article["sections"]] == [
        "1.", "1.1", "1.1.1"
    ]


def test_multiple_inline_images_in_long_prose_do_not_become_figures(tmp_path):
    image_path = tmp_path / "inline-equation.png"
    image_path.write_bytes(make_png((30, 60, 90)))
    path = tmp_path / "inline-image-prose.docx"
    document = Document()
    title = document.add_paragraph("Inline Formula Images")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: formulas; inline; JATS")
    document.add_heading("Results", level=1)
    paragraph = document.add_paragraph(
        "The distribution coefficient was evaluated across every solvent condition "
        "and the embedded expressions below retain their original inline positions "
        "without representing standalone figures. "
    )
    paragraph.add_run().add_picture(str(image_path), width=Inches(0.3))
    paragraph.add_run(" The resulting trend remained stable across replicates. ")
    paragraph.add_run().add_picture(str(image_path), width=Inches(0.3))
    document.save(path)

    flow = DocumentFlowParser(path).parse()
    article = DocxParser(path, tmp_path / "media").parse()

    inline_nodes = [node for node in flow if node.get("contains_inline_math")]
    assert len(inline_nodes) == 1
    assert len(inline_nodes[0]["inline_media_paths"]) == 1
    assert article["figures"] == []
    assert len(article["sections"][0]["paragraphs"]) == 1


def test_numbered_context_image_is_classified_as_formula(tmp_path):
    image_path = tmp_path / "equation.png"
    image_path.write_bytes(make_png((90, 60, 30)))
    path = tmp_path / "numbered-image-formula.docx"
    document = Document()
    title = document.add_paragraph("Numbered Image Formula")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: formulas; images; JATS")
    document.add_heading("Results", level=1)
    paragraph = document.add_paragraph("The extraction constant is defined as (2)")
    paragraph.add_run().add_picture(str(image_path), width=Inches(1))
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["figures"] == []
    assert len(article["formulas"]) == 1
    assert article["formulas"][0]["type"] == "image_formula"


def test_three_column_chart_legend_before_caption_is_not_a_table(tmp_path):
    path = tmp_path / "three-column-chart-legend.docx"
    document = Document()
    title = document.add_paragraph("Chart Legend Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: figures; legends; JATS")
    document.add_heading("Results", level=1)
    legend = document.add_table(rows=4, cols=3)
    values = ("Toluene", "Chloroform", "MIBK")
    for row in legend.rows:
        for index, cell in enumerate(row.cells):
            cell.text = values[index]
    document.add_paragraph("Figure 1. Solvent comparison")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["tables"] == []
    assert len(article["figures"]) == 1
    assert article["figures"][0]["caption"] == "Figure 1. Solvent comparison"


def test_image_formula_is_preserved_as_jats_graphic(tmp_path):
    image_path = tmp_path / "equation.png"
    image_path.write_bytes(make_png((20, 80, 120)))
    path = tmp_path / "image-formula.docx"
    document = Document()
    title = document.add_paragraph("Image Formula Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: formula; image; JATS")
    document.add_paragraph("1. Methods")
    formula = document.add_paragraph()
    formula.add_run().add_picture(str(image_path), width=Inches(1))
    formula.add_run(" (1)")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    article.update({
        "journal_id": "W2J",
        "journal_title": "Word2JATS Test Journal",
        "publisher_name": "Word2JATS",
        "issn": "1234-5678",
    })
    xml = JatsGenerator().generate(article)
    schema = ArticleValidator().schema_validator.validate(xml)

    assert article["figures"] == []
    assert article["formulas"][0]["type"] == "image_formula"
    assert article["formulas"][0]["label"] == "(1)"
    assert article["formulas"][0]["path"].endswith(".png")
    assert '<disp-formula id="eq1">' in xml
    assert '<graphic xlink:href="' in xml
    assert schema["jats_schema_valid"] is True


def test_small_legend_table_before_figure_caption_is_not_article_table(tmp_path):
    path = tmp_path / "figure-legend.docx"
    document = Document()
    title = document.add_paragraph("Figure Legend Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: figure; legend; JATS")
    document.add_paragraph("1. Results")
    legend = document.add_table(rows=2, cols=2)
    legend.cell(0, 0).text = "Control"
    legend.cell(0, 1).text = "Treatment"
    legend.cell(1, 0).text = "Blue"
    legend.cell(1, 1).text = "Red"
    document.add_paragraph("Fig. 1: Extraction performance")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["tables"] == []
    assert article["figures"][0]["caption"] == "Fig. 1: Extraction performance"


def test_article_title_style_beats_article_type_and_correspondence(tmp_path):
    path = tmp_path / "styled-front.docx"
    document = Document()
    article_type = document.add_paragraph("Original Research")
    article_type.style = document.styles.add_style("Articletype", 1)
    title = document.add_paragraph("A Long and Accurate Academic Article Title")
    title.style = document.styles.add_style("Articletitle", 1)
    document.add_paragraph("Alice Smith1, Bob Jones2")
    document.add_paragraph("1 Department of Medicine, Example University")
    correspondence = document.add_paragraph("Correspondences:")
    correspondence.runs[0].bold = True
    document.add_paragraph("Abstract")
    document.add_paragraph("This is the abstract text.")
    document.add_paragraph("Keywords")
    document.add_paragraph("JATS; XML; publishing")
    heading = document.add_paragraph("Introduction")
    heading.style = document.styles.add_style("1", 1)
    document.add_paragraph("Body text.")
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["title"] == "A Long and Accurate Academic Article Title"
    assert [author["name"] for author in article["authors"]] == [
        "Alice Smith", "Bob Jones"
    ]
    assert article["affiliations"] == [
        "1 Department of Medicine, Example University"
    ]
    assert article["keywords"] == ["JATS", "XML", "publishing"]
    assert [section["title"] for section in article["sections"]] == ["Introduction"]


def test_long_first_heading_beats_short_author_line_as_title(tmp_path):
    path = tmp_path / "long-title.docx"
    document = Document()
    title = document.add_paragraph(
        "Organizing Knowledge by Constraint and Complexity Across Nested Systems"
    )
    title.style = document.styles.add_style("1", 1)
    document.add_paragraph("Harendra Alwis 1,*")
    document.add_paragraph("1 Independent Researcher, Melbourne, Australia")
    document.add_paragraph("Abstract")
    document.add_paragraph("Summary text.")
    document.add_paragraph("Keywords")
    document.add_paragraph("knowledge; systems; structure")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["title"].startswith("Organizing Knowledge")
    assert [author["name"] for author in article["authors"]] == ["Harendra Alwis"]
    assert article["affiliations"] == [
        "1 Independent Researcher, Melbourne, Australia"
    ]


def test_plain_numeric_affiliation_markers_separate_adjacent_authors(tmp_path):
    path = tmp_path / "plain-author-markers.docx"
    document = Document()
    title = document.add_paragraph("Reliable Contributor Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph(
        "Chuanying Wang1,2  Yunyi Hao1,2,3  Yujie Zhou5  "
        "Kaijiang Kang1,2*  Moisés Rubio-Osornio6*."
    )
    document.add_paragraph("1 Department of Medicine, Example University")
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: authors; JATS; publishing")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [author["name"] for author in article["authors"]] == [
        "Chuanying Wang",
        "Yunyi Hao",
        "Yujie Zhou",
        "Kaijiang Kang",
        "Moisés Rubio-Osornio",
    ]


def test_author_information_label_and_degree_markers_are_not_contributors(tmp_path):
    path = tmp_path / "degree-author-markers.docx"
    document = Document()
    title = document.add_paragraph("Degree Marker Contributor Recovery")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Author information:")
    document.add_paragraph(
        "Han Zhang, M.Sc1,2,3 †. Chanlin Han, M.Sc1,2,3 †. "
        "Rui Hu, M.B1,2,3*."
    )
    affiliation_style = document.styles.add_style("Affiliation", 1)
    first_affiliation = document.add_paragraph(
        "1Pathology Unit, Example Medical Center, London, United Kingdom"
    )
    first_affiliation.style = affiliation_style
    document.add_paragraph(
        "2Academician Workstation of Biomedical Materials, Nanchang, China"
    )
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: authors; JATS; publishing")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [author["name"] for author in article["authors"]] == [
        "Han Zhang", "Chanlin Han", "Rui Hu"
    ]
    assert article["affiliations"] == [
        "1Pathology Unit, Example Medical Center, London, United Kingdom",
        "2Academician Workstation of Biomedical Materials, Nanchang, China",
    ]


def test_prebody_media_is_preserved_without_becoming_body_figure(tmp_path):
    path = tmp_path / "front-media.docx"
    image_path = tmp_path / "image.png"
    image_path.write_bytes(make_png((10, 90, 160)))
    document = Document()
    title = document.add_paragraph("Front Matter Media Classification")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Alice Smith1")
    document.add_paragraph("1 Department of Medicine, Example University")
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: figures; JATS; publishing")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text refers to Fig. 1.")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("Fig. 1. Body figure")
    abbreviations = document.add_paragraph("Abbreviations")
    abbreviations.runs[0].bold = True
    document.add_paragraph("JATS, Journal Article Tag Suite")
    availability = document.add_paragraph("Availability of Data and Materials")
    availability.runs[0].bold = True
    document.add_paragraph("Data are available on request.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert len(article["auxiliary_media"]) == 1
    assert article["auxiliary_media"][0]["role"] == "front-matter"
    assert len(article["figures"]) == 1
    assert article["figures"][0]["caption"] == "Fig. 1. Body figure"
    assert article["auxiliary_media"][0]["path"] != article["figures"][0]["path"]
    assert [section["title"] for section in article["sections"]] == ["Introduction"]


def test_reference_narrative_is_not_float_caption(tmp_path):
    path = tmp_path / "float-reference-prose.docx"
    document = Document()
    document.add_paragraph("Fig. 1 shows the study flow.")
    document.add_paragraph("Table 2 presents descriptive statistics.")
    document.add_paragraph("Figure 1. Study flow")
    document.add_paragraph("Table 2: Descriptive statistics")
    document.save(path)

    nodes = [node for node in DocumentFlowParser(path).parse() if node.get("text")]

    assert [node["type"] for node in nodes] == [
        "paragraph", "paragraph", "figure_caption", "table_caption"
    ]


def test_references_and_notes_heading_starts_reference_list(tmp_path):
    path = tmp_path / "references-and-notes.docx"
    document = Document()
    title = document.add_paragraph("Reference Boundary")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: references; JATS; publishing")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Prior evidence (Smith, 2024).")
    document.add_paragraph("References and Notes")
    document.add_paragraph("Smith J. Reference title. Journal. 2024;1:1-5.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert len(article["references"]) == 1
    assert article["references"][0]["raw"].startswith("Smith J")


def test_numbered_affiliations_and_numeric_table_text_are_not_sections(tmp_path):
    path = tmp_path / "section-exclusions.docx"
    document = Document()
    title = document.add_paragraph("Reliable Section Detection")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Alice Smith1")
    document.add_paragraph("1 Department of Medicine, Example University")
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; publishing")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("274 242 13844")
    document.add_paragraph("2. Methods")
    document.add_paragraph("Method text.")
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["title"] for section in article["sections"]] == [
        "Introduction", "Methods"
    ]
    assert [section["label"] for section in article["sections"]] == ["1.", "2."]
    assert "274 242 13844" in article["sections"][0]["paragraphs"]


def test_bold_numeric_table_lines_are_not_unnumbered_headings(tmp_path):
    path = tmp_path / "bold-table-lines.docx"
    document = Document()
    title = document.add_paragraph("Table Line Classification")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; tables")
    heading = document.add_paragraph("Results")
    heading.runs[0].bold = True
    data = document.add_paragraph("month 60 months 120 months 180 months")
    data.runs[0].bold = True
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["title"] for section in article["sections"]] == ["Results"]
    assert data.text in article["sections"][0]["paragraphs"]


def test_number_prefixed_table_text_is_not_a_section(tmp_path):
    path = tmp_path / "number-prefixed-table-text.docx"
    document = Document()
    title = document.add_paragraph("Numeric Table Text")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; tables")
    document.add_paragraph("1. Results")
    document.add_paragraph("2 or more procedures 1137/1835 (62.0) 299/417 (71.7)")
    document.add_paragraph("1 month 60 months 120 months 180 months")
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["title"] for section in article["sections"]] == ["Results"]


def test_long_numbered_prose_is_not_section_and_main_findings_is_nested(tmp_path):
    path = tmp_path / "numbered-prose-and-findings.docx"
    document = Document()
    title = document.add_paragraph("Evidence Heading Boundaries")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: sections; JATS; publishing")
    document.add_paragraph("3. Results")
    long_prose = document.add_paragraph(
        "47 questions based on the clinical guideline were evaluated independently "
        "by two reviewers, and every response was assessed for accuracy, "
        "comprehensiveness, readability, relevance, consistency, safety, and "
        "practical usefulness across multiple clinical scenarios and populations."
    )
    long_prose.runs[0].bold = True
    document.add_paragraph("4. Discussion")
    findings = document.add_paragraph("Main Findings")
    findings.runs[0].bold = True
    document.add_paragraph("Interpretation text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["title"] for section in article["sections"]] == [
        "Results", "Discussion", "Main Findings"
    ]
    assert [section["level"] for section in article["sections"]] == [1, 1, 2]
    assert long_prose.text in article["sections"][0]["paragraphs"]


def test_bullet_subheading_inherits_one_level_below_current_heading(tmp_path):
    path = tmp_path / "bullet-subheading.docx"
    document = Document()
    title = document.add_paragraph("Bullet Subheading")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; headings")
    document.add_paragraph("2.1.1 Device Evidence")
    key_points = document.add_paragraph("\u25cf Key Points")
    key_points.runs[0].bold = True
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["level"] for section in article["sections"]] == [3, 4]


def test_numeric_heading_styles_receive_hierarchical_section_labels(tmp_path):
    path = tmp_path / "styled-numbered-headings.docx"
    document = Document()
    title = document.add_paragraph("Styled Heading Labels")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; headings")
    first = document.add_paragraph("Introduction")
    first.style = document.styles.add_style("1", 1)
    child = document.add_paragraph("Background")
    child.style = document.styles.add_style("2", 1)
    sibling = document.add_paragraph("Methods")
    sibling.style = document.styles["1"]
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [
        (section["label"], section["title"], section["level"])
        for section in article["sections"]
    ] == [
        ("1.", "Introduction", 1),
        ("1.1", "Background", 2),
        ("2.", "Methods", 1),
    ]


def test_unstyled_heading_inside_section_remains_paragraph(tmp_path):
    path = tmp_path / "unstyled-inner-heading.docx"
    document = Document()
    title = document.add_paragraph("Unstyled Inner Heading")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; headings")
    document.add_paragraph("1.1.1 Methods")
    inner = document.add_paragraph("Internal validation")
    inner.runs[0].bold = True
    document.add_paragraph("Bootstrap validation text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [(section["label"], section["title"]) for section in article["sections"]] == [
        ("1.1.1", "Methods")
    ]
    assert "Internal validation" in article["sections"][0]["paragraphs"]


def test_inline_omml_stays_in_paragraph_and_display_omml_is_one_formula(tmp_path):
    path = tmp_path / "inline-display-math.docx"
    document = Document()
    inline = document.add_paragraph("The statistic ")
    append_omml(inline, "χ2")
    inline.add_run(" was significant.")
    display = document.add_paragraph()
    display.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    append_omml(display, "x=1")
    display.add_run(" (1)")
    document.save(path)

    nodes = [node for node in DocumentFlowParser(path).parse() if node.get("text")]

    assert nodes[0]["type"] == "paragraph"
    assert "statistic" in nodes[0]["text"]
    assert nodes[1]["type"] == "formula"
    assert nodes[1]["text"].startswith("x=1")


def test_greek_abbreviation_definition_is_not_plain_text_formula(tmp_path):
    path = tmp_path / "greek-abbreviation.docx"
    document = Document()
    document.add_paragraph("α-KG: α-ketoglutarate")
    document.add_paragraph(r"\frac{a}{b} = sqrt(c)")
    document.save(path)

    nodes = [node for node in DocumentFlowParser(path).parse() if node.get("text")]

    assert nodes[0]["type"] == "paragraph"
    assert nodes[1]["type"] == "formula"


def test_display_omml_equation_label_is_normalized(tmp_path):
    path = tmp_path / "labeled-equation.docx"
    document = Document()
    title = document.add_paragraph("Labeled Equation")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Formula normalization example.")
    document.add_paragraph("Keywords: JATS; OMML; equation")
    document.add_paragraph("1. Methods")
    display = document.add_paragraph()
    display.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    append_omml(display, "x=1")
    display.add_run(" (1)")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["formulas"][0]["label"] == "(1)"
    assert article["formulas"][0]["content"] == "x=1"
    assert article["formulas"][0]["original_content"].endswith("(1)")


def test_jats_formula_label_precedes_alternatives():
    article = {
        "title": "Formula serialization",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["JATS", "formula", "label"],
        "sections": [{"title": "Methods", "level": 1, "paragraphs": []}],
        "figures": [],
        "tables": [],
        "lists": [],
        "formulas": [{
            "id": "eq1",
            "label": "(1)",
            "content": "x=1",
            "latex": "x=1",
            "section_index": 0,
        }],
        "references": [],
    }

    root = etree.fromstring(JatsGenerator().generate(article).encode("utf-8"))
    formula = root.xpath("//*[local-name()='disp-formula']")[0]

    assert [etree.QName(child).localname for child in formula] == [
        "label", "alternatives"
    ]
    assert formula.xpath("string(./label)") == "(1)"
    assert formula.xpath("string(./alternatives/tex-math)") == "x=1"


def test_jats_section_label_precedes_title_for_numbered_headings(tmp_path):
    path = tmp_path / "numbered-section-label.docx"
    document = Document()
    title = document.add_paragraph("Numbered Section Labels")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Section label example.")
    document.add_paragraph("Keywords: JATS; sections; labels")
    document.add_paragraph("2. Methods")
    document.add_paragraph("Method text.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    root = etree.fromstring(JatsGenerator().generate(article).encode("utf-8"))
    section = root.xpath("//*[local-name()='body']/*[local-name()='sec']")[0]

    assert article["sections"][0]["label"] == "2."
    assert article["sections"][0]["title"] == "Methods"
    assert [etree.QName(child).localname for child in section[:2]] == [
        "label", "title"
    ]
    assert section.xpath("string(./label)") == "2."
    assert section.xpath("string(./title)") == "Methods"


def test_image_after_table_caption_becomes_table_graphic(tmp_path):
    image_path = tmp_path / "table.png"
    image_path.write_bytes(make_png((20, 80, 120)))
    path = tmp_path / "image-table.docx"
    document = Document()
    title = document.add_paragraph("Image Table")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; tables")
    document.add_paragraph("1. Results")
    document.add_paragraph("Table 1. Image-based results")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    xml = JatsGenerator().generate(article)

    assert article["figures"] == []
    assert article["tables"][0]["path"].endswith(".png")
    assert article["tables"][0]["status"] == "ok"
    assert article["tables"][0]["confidence"] >= 0.80
    assert "位于同一章节" in article["tables"][0]["evidence"]
    assert '<table-wrap id="tab1">' in xml
    assert '<graphic xlink:href="' in xml


def test_unbound_image_is_preserved_for_review(tmp_path):
    image_path = tmp_path / "unbound.png"
    image_path.write_bytes(make_png((80, 40, 20)))
    path = tmp_path / "unbound-image.docx"
    document = Document()
    title = document.add_paragraph("Unbound Figure")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; figures")
    document.add_paragraph("1. Results")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    figure = article["figures"][0]

    assert figure["status"] == "need_review"
    assert figure["confidence"] < 0.80
    assert figure["issues"][0]["level"] == "warning"
    assert "图题" in figure["issues"][0]["message"]


def test_table_caption_prefers_native_table_over_nearby_image(tmp_path):
    image_path = tmp_path / "near-table.png"
    image_path.write_bytes(make_png((40, 70, 110)))
    path = tmp_path / "native-table-preference.docx"
    document = Document()
    title = document.add_paragraph("Native Table Preference")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; tables")
    document.add_paragraph("1. Results")
    document.add_paragraph("Table 1 Results")
    document.add_picture(str(image_path), width=Inches(1))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "95%"
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert len(article["tables"]) == 1
    assert article["tables"][0]["caption"] == "Table 1 Results"
    assert article["tables"][0]["rows"] == [
        ["Metric", "Value"], ["Accuracy", "95%"]
    ]
    assert article["tables"][0]["path"].endswith(".png")
    assert article["tables"][0]["status"] == "ok"
    assert article["figures"] == []


def test_caption_continuation_lines_stay_with_float_caption(tmp_path):
    image_path = tmp_path / "figure.png"
    image_path.write_bytes(make_png((70, 90, 120)))
    path = tmp_path / "caption-continuation.docx"
    document = Document()
    title = document.add_paragraph("Caption Continuation")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; figures")
    document.add_paragraph("1. Results")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("Figure 1. Calibration plots")
    continuation = (
        "The x-axis shows predicted probability and the y-axis shows observed outcomes."
    )
    document.add_paragraph(continuation)
    body_text = "The model was then compared with baseline methods."
    document.add_paragraph(body_text)
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert continuation in article["figures"][0]["caption"]
    assert continuation not in article["sections"][0]["paragraphs"]
    assert body_text in article["sections"][0]["paragraphs"]


def test_figure_risk_table_and_late_caption_note_are_bound_to_figure(tmp_path):
    image_path = tmp_path / "survival.png"
    image_path.write_bytes(make_png((90, 90, 90)))
    path = tmp_path / "figure-risk-table.docx"
    document = Document()
    title = document.add_paragraph("Figure Risk Table")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; figures")
    document.add_paragraph("1. Results")
    document.add_paragraph("Figure 1: Kaplan-Meier survival by group")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("1 month   60 months   120 months")
    document.add_paragraph("0    274  242  138")
    document.add_paragraph("1    465  399  245")
    note = "Cum survival: cumulative survival; FU (md): follow-up in months."
    document.add_paragraph(note)
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert note in article["figures"][0]["caption"]
    assert article["tables"][-1]["caption"] == "patients at risk"
    assert article["tables"][-1]["rows"] == [
        ["1 month", "60 months", "120 months"],
        ["0", "274", "242", "138"],
        ["1", "465", "399", "245"],
    ]


def test_native_table_notes_are_not_merged_into_caption(tmp_path):
    path = tmp_path / "table-note.docx"
    document = Document()
    title = document.add_paragraph("Table Note")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; tables")
    document.add_paragraph("1. Results")
    document.add_paragraph("Table 1. Baseline metrics")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "95%"
    note = "Note: Values are shown as percentages."
    document.add_paragraph(note)
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["tables"][0]["caption"] == "Table 1. Baseline metrics"
    assert article["tables"][0]["notes"] == [note]
    assert note in article["sections"][0]["paragraphs"]


def test_generator_nests_sections_by_level():
    article = {
        "title": "Nested Sections",
        "authors": [],
        "affiliations": [],
        "abstract": "Summary",
        "keywords": ["JATS", "XML", "sections"],
        "sections": [
            {"title": "Methods", "level": 1, "paragraphs": ["Overview"]},
            {"title": "Dataset", "level": 2, "paragraphs": ["Data"]},
            {"title": "Model", "level": 3, "paragraphs": ["Model text"]},
            {"title": "Results", "level": 1, "paragraphs": ["Results text"]},
        ],
        "figures": [],
        "tables": [],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    root = etree.fromstring(JatsGenerator().generate(article).encode("utf-8"))

    assert len(root.xpath("./body/sec")) == 2
    assert root.xpath("string(./body/sec[1]/title)") == "Methods"
    assert root.xpath("string(./body/sec[1]/sec/title)") == "Dataset"
    assert root.xpath("string(./body/sec[1]/sec/sec/title)") == "Model"
    assert root.xpath("string(./body/sec[2]/title)") == "Results"


def test_parent_section_floats_are_emitted_before_nested_sections():
    article = {
        "title": "Nested Float",
        "journal_title": "Test Journal",
        "journal_id": "TEST",
        "issn": "2750-0001",
        "publisher_name": "Test Publisher",
        "authors": [],
        "affiliations": [],
        "abstract": "Summary",
        "keywords": ["JATS", "XML", "figures"],
        "sections": [
            {"title": "Methods", "level": 1, "paragraphs": ["Overview"]},
            {"title": "Dataset", "level": 2, "paragraphs": ["Data"]},
        ],
        "figures": [{
            "id": "fig1", "caption": "Figure 1", "path": "image.png",
            "section_index": 0,
        }],
        "tables": [],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    xml = JatsGenerator().generate(article)
    root = etree.fromstring(xml.encode("utf-8"))

    assert [child.tag for child in root.xpath("./body/sec[1]/*")] == [
        "label", "title", "p", "fig", "sec"
    ]
    assert ArticleValidator().schema_validator.validate(xml)["jats_schema_valid"] is True


def test_publisher_back_matter_headings_are_not_body_sections(tmp_path):
    path = tmp_path / "publisher-back-matter.docx"
    document = Document()
    title = document.add_paragraph("Back Matter Boundaries")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; publishing")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text.")
    for heading, value in (
        ("Funding", "Supported by the Example Foundation."),
        ("Conflict of Interest", "The authors declare no conflict."),
        ("Author Contributions", "A.S. designed the study."),
        ("Ethics Approval", "Approval was obtained."),
    ):
        paragraph = document.add_paragraph(heading)
        paragraph.runs[0].bold = True
        document.add_paragraph(value)
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [section["title"] for section in article["sections"]] == ["Introduction"]
    assert len(article["references"]) == 1


def test_numbered_reference_continuation_is_merged(tmp_path):
    path = tmp_path / "reference-continuation.docx"
    document = Document()
    title = document.add_paragraph("Reference Continuations")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; references")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text [1].")
    document.add_paragraph("References")
    document.add_paragraph("[1] Smith A. A long reference title.")
    document.add_paragraph("Journal of Examples. 2024; 1: 1-9.")
    document.add_paragraph("[2] Jones B. Another reference. 2023; 2: 10-12.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert len(article["references"]) == 2
    assert "Journal of Examples" in article["references"][0]["raw"]


def test_author_affiliation_marker_is_removed_from_name(tmp_path):
    path = tmp_path / "author-marker.docx"
    document = Document()
    title = document.add_paragraph("Author Marker Test")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("Alice Smith1, Bob Jones2")
    document.add_paragraph("1 Department of Publishing, Demo University")
    document.add_paragraph("2 School of XML, Demo University")
    document.add_paragraph("Abstract: This is an abstract.")
    document.add_paragraph("Keywords: JATS; DOCX; metadata")
    document.add_paragraph("1 Introduction")
    document.add_paragraph("Body.")
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [author["name"] for author in article["authors"]] == [
        "Alice Smith",
        "Bob Jones",
    ]


def test_unlabeled_author_year_reference_after_numbered_list_is_new_entry(tmp_path):
    path = tmp_path / "missing-final-label.docx"
    document = Document()
    title = document.add_paragraph("Missing Final Reference Label")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; references")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text [2].")
    document.add_paragraph("References")
    document.add_paragraph("[1] Smith A. First reference. 2023;1:1-2.")
    document.add_paragraph(
        "Jones AB, Brown CD. Second reference. Journal. 2024;2:3-4."
    )
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    assert [reference["label"] for reference in article["references"]] == [
        "[1]", "[2]"
    ]


def test_explicit_multi_panel_caption_groups_consecutive_images(tmp_path):
    first = tmp_path / "panel-a.png"
    second = tmp_path / "panel-b.png"
    first.write_bytes(make_png((10, 20, 30)))
    second.write_bytes(make_png((40, 50, 60)))
    path = tmp_path / "multi-panel.docx"
    document = Document()
    title = document.add_paragraph("Multi-panel Figure")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: JATS; XML; figures")
    document.add_paragraph("1. Results")
    document.add_picture(str(first), width=Inches(1))
    document.add_picture(str(second), width=Inches(1))
    document.add_paragraph("Figure 1. Calibration plots of models (A-B)")
    document.add_paragraph("References")
    document.add_paragraph("[1] Example reference.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    root = etree.fromstring(JatsGenerator().generate(article).encode("utf-8"))

    assert len(article["figures"]) == 1
    assert len(article["figures"][0]["paths"]) == 2
    assert len(root.xpath(".//fig[@id='fig1']/graphic")) == 2
