from docx import Document
from docx.oxml import parse_xml
from lxml import etree

from app.services.docx_parser import DocxParser
from app.services.jats_generator import JatsGenerator
from app.services.omml_converter import OmmlConverter
from app.services.validator import ArticleValidator
from app.services.quality_scorer import QualityScorer
from app.services.jats_schema_validator import JatsSchemaValidator


OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MATHML_NS = "http://www.w3.org/1998/Math/MathML"

COMPLEX_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}">
  <m:f><m:num><m:r><m:t>a</m:t></m:r></m:num><m:den><m:r><m:t>b</m:t></m:r></m:den></m:f>
  <m:r><m:t>+</m:t></m:r>
  <m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>
  <m:r><m:t>+</m:t></m:r>
  <m:sSub><m:e><m:r><m:t>y</m:t></m:r></m:e><m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub>
  <m:r><m:t>+</m:t></m:r>
  <m:rad><m:e><m:r><m:t>z</m:t></m:r></m:e></m:rad>
  <m:r><m:t>+</m:t></m:r>
  <m:nary>
    <m:naryPr><m:chr m:val="∑"/></m:naryPr>
    <m:sub><m:r><m:t>i=1</m:t></m:r></m:sub>
    <m:sup><m:r><m:t>n</m:t></m:r></m:sup>
    <m:e><m:r><m:t>x</m:t></m:r></m:e>
  </m:nary>
  <m:r><m:t>+</m:t></m:r>
  <m:d><m:dPr><m:begChr m:val="("/><m:endChr m:val=")"/></m:dPr><m:e><m:r><m:t>q</m:t></m:r></m:e></m:d>
</m:oMath>
""".strip()

MATRIX_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}"><m:m>
  <m:mr><m:e><m:r><m:t>a</m:t></m:r></m:e><m:e><m:r><m:t>b</m:t></m:r></m:e></m:mr>
  <m:mr><m:e><m:r><m:t>c</m:t></m:r></m:e><m:e><m:r><m:t>d</m:t></m:r></m:e></m:mr>
</m:m></m:oMath>
""".strip()

EQARR_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}"><m:eqArr>
  <m:e><m:r><m:t>x=1</m:t></m:r></m:e>
  <m:e><m:r><m:t>y=2</m:t></m:r></m:e>
  <m:e><m:r><m:t>z=3</m:t></m:r></m:e>
</m:eqArr></m:oMath>
""".strip()

CASES_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}"><m:d>
  <m:dPr><m:begChr m:val="{{"/><m:endChr m:val=""/></m:dPr>
  <m:e><m:eqArr>
    <m:e><m:r><m:t>x, x&gt;0</m:t></m:r></m:e>
    <m:e><m:r><m:t>-x, x≤0</m:t></m:r></m:e>
  </m:eqArr></m:e>
</m:d></m:oMath>
""".strip()

NARY_LIMITS_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}">
  <m:nary><m:naryPr><m:chr m:val="∑"/></m:naryPr><m:sub><m:r><m:t>i=1</m:t></m:r></m:sub><m:sup><m:r><m:t>n</m:t></m:r></m:sup><m:e><m:r><m:t>x_i</m:t></m:r></m:e></m:nary>
  <m:nary><m:naryPr><m:chr m:val="∫"/></m:naryPr><m:sub><m:r><m:t>0</m:t></m:r></m:sub><m:sup><m:r><m:t>1</m:t></m:r></m:sup><m:e><m:r><m:t>f(x)dx</m:t></m:r></m:e></m:nary>
</m:oMath>
""".strip()

ACCENT_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}">
  <m:acc><m:accPr><m:chr m:val="^"/></m:accPr><m:e><m:r><m:t>x</m:t></m:r></m:e></m:acc>
  <m:acc><m:accPr><m:chr m:val="¯"/></m:accPr><m:e><m:r><m:t>y</m:t></m:r></m:e></m:acc>
  <m:acc><m:accPr><m:chr m:val="˙"/></m:accPr><m:e><m:r><m:t>z</m:t></m:r></m:e></m:acc>
  <m:acc><m:accPr><m:chr m:val="~"/></m:accPr><m:e><m:r><m:t>q</m:t></m:r></m:e></m:acc>
</m:oMath>
""".strip()

PARTIAL_OMML = f"""
<m:oMath xmlns:m="{OMML_NS}">
  <m:acc><m:accPr><m:chr m:val="⏞"/></m:accPr><m:e><m:r><m:t>x</m:t></m:r></m:e></m:acc>
  <m:unknownComplex><m:r><m:t>kept</m:t></m:r></m:unknownComplex>
</m:oMath>
""".strip()


def add_omml(document: Document, omml: str) -> None:
    paragraph = document.add_paragraph()
    paragraph._p.append(parse_xml(omml))


def test_omml_converter_supports_common_presentation_mathml_structures():
    result = OmmlConverter().convert(COMPLEX_OMML)
    math = etree.fromstring(result["mathml"].encode("utf-8"))

    assert math.tag == f"{{{MATHML_NS}}}math"
    assert math.xpath("count(.//mml:mfrac)", namespaces={"mml": MATHML_NS}) == 1
    assert math.xpath("count(.//mml:msup)", namespaces={"mml": MATHML_NS}) >= 1
    assert math.xpath("count(.//mml:msub)", namespaces={"mml": MATHML_NS}) == 1
    assert math.xpath("count(.//mml:msqrt)", namespaces={"mml": MATHML_NS}) == 1
    assert math.xpath("count(.//mml:munderover)", namespaces={"mml": MATHML_NS}) == 1
    assert "(" in result["latex"] and ")" in result["latex"]
    assert r"\frac{a}{b}" in result["latex"]
    assert r"\sqrt{z}" in result["latex"]
    assert r"\sum_{i=1}^{n}" in result["latex"]


def test_docx_parser_extracts_raw_omml_mathml_and_latex(tmp_path):
    path = tmp_path / "native-formula.docx"
    document = Document()
    document.add_paragraph("原生公式测试")
    document.add_paragraph("摘要：测试 OMML 转换。")
    document.add_paragraph("关键词：OMML；MathML；JATS")
    document.add_paragraph("1 方法")
    add_omml(document, COMPLEX_OMML)
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()
    formula = article["formulas"][0]

    assert formula["type"] == "omml"
    assert formula["omml"].startswith("<m:oMath")
    assert "<mml:math" in formula["mathml"]
    assert r"\frac{a}{b}" in formula["latex"]
    assert formula["section_index"] == 0


def test_generator_outputs_mathml_and_tex_alternatives():
    converted = OmmlConverter().convert(COMPLEX_OMML)
    article = {
        "title": "公式测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["OMML", "MathML", "JATS"],
        "sections": [{"title": "方法", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "tables": [],
        "lists": [],
        "formulas": [
            {
                "id": "eq1",
                "content": "a/b",
                "omml": COMPLEX_OMML,
                "mathml": converted["mathml"],
                "latex": converted["latex"],
                "type": "omml",
                "section_index": 0,
            }
        ],
        "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert '<disp-formula id="eq1">' in xml
    assert "<alternatives>" in xml
    assert "<mml:math" in xml
    assert "<mml:mfrac>" in xml
    assert "<![CDATA[\\frac{a}{b}" in xml


def test_validator_warns_when_omml_mathml_conversion_is_unavailable():
    article = {
        "title": "公式回退测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["OMML", "MathML", "测试"],
        "sections": [{"title": "方法", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "tables": [],
        "lists": [],
        "formulas": [
            {
                "id": "eq1",
                "content": "unsupported",
                "omml": "<m:oMath/>",
                "mathml": "",
                "latex": "",
                "type": "omml",
                "section_index": 0,
            }
        ],
        "references": [],
    }

    result = ArticleValidator().validate(article, JatsGenerator().generate(article))

    assert "公式 eq1 的 OMML 无法转换为 MathML，已保留文本回退。" in result["warnings"]


def test_omml_matrix_to_mathml():
    result = OmmlConverter().convert(MATRIX_OMML)
    assert result["conversion_status"] == "success"
    assert "matrix" in result["supported_features"]
    assert result["mathml"].count("<mml:mtr>") == 2
    assert result["mathml"].count("<mml:mtd>") == 4
    assert r"\begin{matrix}" in result["latex"]


def test_omml_equation_array_to_mathml():
    result = OmmlConverter().convert(EQARR_OMML)
    assert result["conversion_status"] == "success"
    assert "equation_array" in result["supported_features"]
    assert result["mathml"].count("<mml:mtr>") == 3
    assert r"\begin{aligned}" in result["latex"]


def test_omml_cases_partial_or_success():
    result = OmmlConverter().convert(CASES_OMML)
    assert result["conversion_status"] in {"success", "partial"}
    assert "cases" in result["supported_features"]
    assert 'open="{" close=""' in result["mathml"]
    assert r"\begin{cases}" in result["latex"]


def test_omml_nary_limits():
    result = OmmlConverter().convert(NARY_LIMITS_OMML)
    assert result["conversion_status"] == "success"
    assert result["mathml"].count("<mml:munderover>") == 2
    assert r"\sum_{i=1}^{n}" in result["latex"]
    assert r"\int_{0}^{1}" in result["latex"]


def test_omml_accent():
    result = OmmlConverter().convert(ACCENT_OMML)
    assert result["conversion_status"] == "success"
    assert result["mathml"].count("<mml:mover") == 4
    assert r"\hat{x}" in result["latex"]
    assert r"\bar{y}" in result["latex"]
    assert r"\dot{z}" in result["latex"]
    assert r"\tilde{q}" in result["latex"]


def test_omml_partial_does_not_break_jats_generation():
    converted = OmmlConverter().convert(PARTIAL_OMML)
    formula = {
        "id": "eq1", "content": "x kept", "type": "omml", "omml": PARTIAL_OMML,
        "mathml": converted["mathml"], "latex": converted["latex"],
        "conversion_status": converted["conversion_status"],
        "supported_features": converted["supported_features"],
        "unsupported_features": converted["unsupported_features"],
        "issues": converted["issues"], "section_index": 0,
    }
    article = {
        "title": "Partial", "abstract": "Abstract", "keywords": ["a", "b", "c"],
        "journal_title": "Test Journal", "journal_id": "TEST", "issn": "1234-5678",
        "publisher_name": "Test Publisher",
        "sections": [{"title": "Method", "paragraphs": ["Text"]}],
        "authors": [{"name": "Alice Smith", "orcid": "0000-0002-0000-0001", "affiliation_ids": ["aff1"]}],
        "affiliations": ["Publishing Lab"], "figures": [], "tables": [], "lists": [],
        "formulas": [formula], "references": [],
    }
    xml = JatsGenerator().generate(article)
    validation = ArticleValidator().validate(article, xml)
    quality = QualityScorer().score(article, validation)

    assert converted["conversion_status"] == "partial"
    assert converted["unsupported_features"]
    assert "<disp-formula" in xml and "<mml:math" in xml
    assert any("人工复核" in warning for warning in validation["warnings"])
    assert any(issue["module"] == "formula" for issue in quality["issues"])
    assert JatsSchemaValidator().validate(xml)["jats_schema_valid"] is True
