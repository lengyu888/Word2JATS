from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from PIL import Image
from docx import Document
from fastapi.testclient import TestClient

from app.services.docx_parser import DocxParser
from app.services.visual_preview_builder import VisualPreviewBuilder


def _tiff_bytes() -> bytes:
    output = BytesIO()
    Image.new("RGB", (96, 48), (12, 92, 150)).save(output, format="TIFF")
    return output.getvalue()


def test_tiff_extraction_keeps_original_and_creates_png_preview(tmp_path):
    docx_path = tmp_path / "source.docx"
    with ZipFile(docx_path, "w") as archive:
        archive.writestr("word/media/image1.tif", _tiff_bytes())

    media_dir = tmp_path / "media"
    saved = DocxParser(docx_path, media_dir)._save_flow_image("word/media/image1.tif", 1)

    assert saved["path"].endswith("figure_1.tif")
    assert saved["preview_path"].endswith("figure_1_preview.png")
    assert (media_dir / "figure_1.tif").is_file()
    preview = media_dir / "figure_1_preview.png"
    assert preview.is_file()
    with Image.open(preview) as image:
        assert image.format == "PNG"


def test_tiff_uses_preview_url_but_keeps_original_media_path():
    article = {
        "sections": [],
        "figures": [{
            "id": "fig1",
            "caption": "Figure 1 TIFF source",
            "path": "temp/demo/media/figure_1.tif",
            "preview_path": "temp/demo/media/figure_1_preview.png",
            "section_index": -1,
        }],
        "tables": [],
    }

    VisualPreviewBuilder().enrich(article, "a" * 32, {"issues": []})

    figure = article["figures"][0]
    assert figure["filename"] == "figure_1.tif"
    assert figure["preview_filename"] == "figure_1_preview.png"
    assert figure["media_url"] == ""
    assert figure["preview_url"] == f"/api/media/{'a' * 32}/figure_1_preview.png"
    assert figure["status"] == "ok"


def test_convert_endpoint_returns_png_preview_for_tiff(tmp_path):
    from app.main import app

    source_image = tmp_path / "figure.tif"
    source_image.write_bytes(_tiff_bytes())
    source_docx = tmp_path / "figure.docx"
    document = Document()
    document.add_paragraph("TIFF preview article")
    document.add_paragraph("Abstract: A preview integration test.")
    document.add_paragraph("Keywords: TIFF, preview, JATS")
    document.add_paragraph("1 Results")
    document.add_picture(str(source_image))
    document.add_paragraph("Figure 1. TIFF preview")
    document.save(source_docx)

    response = TestClient(app).post(
        "/api/convert",
        files={
            "file": (
                source_docx.name,
                source_docx.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    figure = response.json()["article"]["figures"][0]
    assert figure["path"].endswith(".tif")
    assert figure["preview_url"].endswith("_preview.png")
    preview = TestClient(app).get(figure["preview_url"])
    assert preview.status_code == 200
    assert preview.headers["content-type"] == "image/png"
    assert preview.content.startswith(b"\x89PNG")
