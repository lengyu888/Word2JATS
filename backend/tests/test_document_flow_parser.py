from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches

from app.services.document_flow_parser import DocumentFlowParser
from app.services.docx_parser import DocxParser
from tests.test_services import make_png


def add_omml_formula(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    math.append(run)
    math_para.append(math)
    paragraph._p.append(math_para)


def build_flow_docx(path) -> None:
    image_path = path.parent / "flow.png"
    image_path.write_bytes(make_png((20, 100, 130)))

    document = Document()
    document.add_paragraph("真实文档流测试")
    document.add_paragraph("摘要：测试真实顺序。")
    document.add_paragraph("关键词：流；图片；表格")
    document.add_paragraph("1 第一节")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("图1 第一节图片")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "准确率"
    table.cell(1, 1).text = "95%"
    document.add_paragraph("表1 第一节表格")
    document.add_paragraph("2 第二节")
    add_omml_formula(document, "x=1")
    document.save(path)


def test_document_flow_parser_preserves_body_order_and_node_types(tmp_path):
    path = tmp_path / "flow.docx"
    build_flow_docx(path)

    nodes = DocumentFlowParser(path).parse()
    relevant = [node for node in nodes if node["type"] != "paragraph"][-7:]

    assert [node["type"] for node in relevant] == [
        "heading",
        "image",
        "figure_caption",
        "table",
        "table_caption",
        "heading",
        "formula",
    ]
    assert relevant[1]["media_path"].startswith("word/media/")
    assert relevant[3]["rows"] == [["指标", "结果"], ["准确率", "95%"]]
    assert relevant[-1]["text"] == "x=1"
    assert relevant[-1]["formula_type"] == "omml"


def test_document_flow_parser_classifies_title_paragraph_and_list(tmp_path):
    path = tmp_path / "paragraph-types.docx"
    document = Document()
    title = document.add_paragraph("标题样式")
    title.style = "Title"
    document.add_paragraph("普通正文")
    document.add_paragraph("（1）列表内容")
    document.save(path)

    nodes = DocumentFlowParser(path).parse()

    assert [node["type"] for node in nodes[:3]] == ["title", "paragraph", "list"]


def test_docx_parser_assigns_flow_media_to_their_real_sections(tmp_path):
    path = tmp_path / "flow.docx"
    build_flow_docx(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["figures"][0]["caption"] == "图1 第一节图片"
    assert article["figures"][0]["section_index"] == 0
    assert article["tables"][0]["caption"] == "表1 第一节表格"
    assert article["tables"][0]["section_index"] == 0
    assert article["formulas"][0]["content"] == "x=1"
    assert article["formulas"][0]["section_index"] == 1


def test_docx_parser_does_not_bind_captions_across_section_boundaries(tmp_path):
    image_path = tmp_path / "section.png"
    image_path.write_bytes(make_png((30, 80, 120)))
    path = tmp_path / "section-boundary.docx"
    document = Document()
    document.add_paragraph("章节边界测试")
    document.add_paragraph("1 第一节")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("2 第二节")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("图2 第二节图片")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["figures"][0]["caption"] == ""
    assert article["figures"][0]["section_index"] == 0
    assert article["figures"][1]["caption"] == "图2 第二节图片"
    assert article["figures"][1]["section_index"] == 1
