from fastapi.testclient import TestClient
from io import BytesIO
import json
from pathlib import Path
import pytest
import zipfile
from docx import Document
from docx.shared import Inches, Pt

from app.main import app
from tests.test_services import build_figure_docx, build_sample_docx, make_png


client = TestClient(app)


def test_health_endpoint():
    response = client.get("/api/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_demo_document_endpoint():
    response = client.get("/api/demo-document")

    assert response.status_code == 200
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers["content-type"]
    assert response.content.startswith(b"PK")


def test_demo_documents_endpoint_lists_and_downloads_both_demo_files():
    response = client.get("/api/demo-documents")

    assert response.status_code == 200
    documents = response.json()["documents"]
    assert [item["filename"] for item in documents] == [
        "official-sample-1-group-1.docx",
        "official-sample-2-group-1.docx",
        "official-sample-3-group-1.docx",
        "official-sample-4-group-1.docx",
        "official-sample-5-group-1.docx",
    ]
    assert documents[0]["official_xml"].endswith(".xml")

    for document in documents:
        download = client.get(document["download_url"])
        assert download.status_code == 200
        assert download.content.startswith(b"PK")


def test_demo_documents_endpoint_rejects_files_outside_allowlist():
    response = client.get("/api/demo-documents/../README.md")

    assert response.status_code in {404, 405}


def test_official_demo_conversion_returns_comparison_report():
    listing = client.get("/api/demo-documents")
    if listing.status_code == 404:
        pytest.skip("official samples are not available in this workspace")
    document = listing.json()["documents"][0]
    download = client.get(document["download_url"])

    response = client.post(
        "/api/convert",
        files={
            "file": (
                document["filename"],
                download.content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    comparison = response.json()["official_comparison"]
    assert comparison["available"] is True
    assert comparison["official_xml"] == document["official_xml"]
    assert 0 <= comparison["similarity_score"] <= 100
    assert "sec" in comparison["counts"]["generated"]
    assert isinstance(comparison["differences"], list)


def test_profiles_endpoint_and_convert_profile_parameter():
    profiles = client.get("/api/profiles")
    assert profiles.status_code == 200
    assert any(item["id"] == "english_journal" for item in profiles.json()["profiles"])

    response = client.post(
        "/api/convert",
        data={"profile": "english_journal"},
        files={"file": ("sample.docx", build_sample_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 200
    assert response.json()["article"]["profile"] == "english_journal"
    assert response.json()["article"]["lang"] == "en"


def test_convert_endpoint_returns_all_outputs():
    response = client.post(
        "/api/convert",
        files={
            "file": (
                "sample.docx",
                build_sample_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is True
    assert payload["article"]["title"] == "面向出版的智能结构化转换"
    assert payload["xml"].startswith("<?xml")
    assert 'dtd-version="1.3"' in payload["xml"]
    assert payload["validation"]["passed"] is True
    assert payload["validation"]["jats_schema_valid"] is True
    assert payload["validation"]["auto_fix"]["attempted"] is False
    assert isinstance(payload["validation"]["auto_fix"]["applied_fixes"], list)
    assert 0 <= payload["quality_report"]["total_score"] <= 100
    assert "metadata_score" in payload["quality_report"]["scores"]
    assert payload["article"]["document_flow_view"]
    assert any(item["jats_tag"] == "article-title" for item in payload["article"]["document_flow_view"])
    assert len(payload["conversion_id"]) == 32
    assert isinstance(payload["media_paths"], list)


def test_convert_figures_include_secure_media_url(tmp_path):
    path = tmp_path / "figure.docx"
    build_figure_docx(path, ["图1 测试图片"], 1)

    response = client.post(
        "/api/convert",
        files={"file": (path.name, path.read_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    payload = response.json()
    figure = payload["article"]["figures"][0]
    assert figure["filename"].endswith(".png")
    assert figure["media_url"] == f"/api/media/{payload['conversion_id']}/{figure['filename']}"
    assert payload["media_paths"] == [figure["path"]]
    media = client.get(figure["media_url"])
    assert media.status_code == 200
    assert media.headers["content-type"] == "image/png"
    assert media.content.startswith(b"\x89PNG")


def test_convert_preserves_prebody_auxiliary_media_for_package_export(tmp_path):
    image_path = tmp_path / "graphical-abstract.png"
    image_path.write_bytes(make_png((30, 100, 170)))
    docx_path = tmp_path / "front-media.docx"
    document = Document()
    title = document.add_paragraph("Auxiliary Front Matter Media")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("Alice Smith1")
    document.add_paragraph("1 Department of Medicine, Example University")
    document.add_paragraph("Abstract: Summary")
    document.add_paragraph("Keywords: media; JATS; publishing")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Body text.")
    document.save(docx_path)

    response = client.post(
        "/api/convert",
        files={
            "file": (
                docx_path.name,
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["article"]["figures"] == []
    auxiliary = payload["article"]["auxiliary_media"]
    assert len(auxiliary) == 1
    assert auxiliary[0]["role"] == "front-matter"
    assert payload["media_paths"] == [auxiliary[0]["path"]]


def test_media_endpoint_rejects_invalid_or_missing_paths():
    from app.routers.convert import TEMP_ROOT
    conversion_id = "b" * 32
    media_dir = TEMP_ROOT / conversion_id / "media"
    media_dir.mkdir(parents=True, exist_ok=True)
    (media_dir / "valid.png").write_bytes(make_png((20, 80, 120)))

    assert client.get(f"/api/media/{conversion_id}/valid.png").status_code == 200
    assert client.get(f"/api/media/{conversion_id}/missing.png").status_code == 404
    assert client.get(f"/api/media/not-a-valid-id/valid.png").status_code in {400, 404}
    assert client.get(f"/api/media/{conversion_id}/..%2Fsecret.png").status_code in {400, 404}


def test_convert_endpoint_rejects_non_docx():
    response = client.post(
        "/api/convert",
        files={"file": ("paper.txt", b"not a docx", "text/plain")},
    )

    assert response.status_code == 400


def test_generate_xml_endpoint_uses_corrected_article():
    response = client.post(
        "/api/generate-xml",
        json={
            "article": {
                "title": "人工校正后的标题",
                "authors": [{"name": "张三", "orcid": ""}],
                "affiliations": ["校正大学"],
                "abstract": "人工校正后的摘要",
                "keywords": ["校正", "JATS"],
                "sections": [
                    {
                        "title": "校正后的章节",
                        "level": 1,
                        "paragraphs": ["校正后的正文段落。"],
                    }
                ],
                "figures": [],
                "lists": [],
                "formulas": [],
                "references": [{"raw": "[1] 校正后的参考文献"}],
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is True
    assert "<article-title>人工校正后的标题</article-title>" in payload["xml"]
    assert "<title>校正后的章节</title>" in payload["xml"]
    assert payload["validation"]["passed"] is True
    assert "remaining_schema_errors" in payload["validation"]["auto_fix"]
    assert payload["quality_report"]["issues"] is not None
    assert payload["article"]["document_flow_view"]


def test_generate_xml_endpoint_normalizes_missing_optional_collections():
    response = client.post(
        "/api/generate-xml",
        json={
            "article": {
                "title": "最小文章",
                "abstract": "摘要",
                "keywords": ["关键词"],
                "sections": [{"title": "引言", "paragraphs": ["正文"]}],
            }
        },
    )

    assert response.status_code == 200
    assert response.json()["validation"]["passed"] is True


def test_generate_xml_endpoint_accepts_legacy_formula_fields():
    response = client.post(
        "/api/generate-xml",
        json={
            "article": {
                "title": "公式兼容测试",
                "abstract": "摘要",
                "keywords": ["公式", "JATS", "测试"],
                "sections": [{"title": "方法", "paragraphs": ["正文"]}],
                "formulas": [{"plain_text": "E = mc²", "section_index": 0}],
            }
        },
    )

    assert response.status_code == 200
    assert '<disp-formula id="eq1">' in response.json()["xml"]
    assert "<![CDATA[E = mc²]]>" in response.json()["xml"]


def test_generate_xml_endpoint_normalizes_legacy_reference_raw():
    response = client.post(
        "/api/generate-xml",
        json={
            "article": {
                "title": "参考文献兼容测试",
                "abstract": "摘要",
                "keywords": ["参考文献", "JATS", "测试"],
                "sections": [{"title": "引言", "paragraphs": ["正文"]}],
                "references": [{"raw": "[1] Legacy citation."}],
            }
        },
    )

    assert response.status_code == 200
    xml = response.json()["xml"]
    assert '<ref id="ref1">' in xml
    assert "<label>[1]</label>" in xml
    assert "<mixed-citation>Legacy citation.</mixed-citation>" in xml


def test_generate_xml_endpoint_accepts_publishing_metadata():
    response = client.post(
        "/api/generate-xml",
        json={
            "article": {
                "title": "元数据测试",
                "doi": "10.1234/test",
                "journal_title": "测试期刊",
                "journal_id": "TEST",
                "publisher_name": "测试出版社",
                "subject": "出版技术",
                "pub_year": "2026",
                "pub_month": "06",
                "pub_day": "10",
                "abstract": "摘要",
                "keywords": ["JATS", "XML", "出版"],
                "sections": [{"title": "引言", "paragraphs": ["正文"]}],
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["validation"]["passed"] is True
    assert "<journal-title>测试期刊</journal-title>" in payload["xml"]
    assert '<article-id pub-id-type="doi">10.1234/test</article-id>' in payload["xml"]


def test_generate_xml_endpoint_accepts_tables():
    response = client.post(
        "/api/generate-xml",
        json={
            "article": {
                "title": "表格接口测试",
                "abstract": "摘要",
                "keywords": ["表格", "JATS", "测试"],
                "sections": [{"title": "结果", "paragraphs": ["正文"]}],
                "tables": [
                    {
                        "id": "tab1",
                        "caption": "表1 实验结果",
                        "rows": [["指标", "结果"], ["准确率", "95%"]],
                        "section_index": 0,
                    }
                ],
            }
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert '<table-wrap id="tab1">' in payload["xml"]


def test_batch_convert_returns_success_and_failure_per_file():
    response = client.post(
        "/api/batch-convert",
        files=[
            (
                "files",
                (
                    "sample.docx",
                    build_sample_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
            ("files", ("invalid.txt", b"not a docx", "text/plain")),
        ],
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["success"] is False
    assert len(payload["results"]) == 2
    assert payload["results"][0]["filename"] == "sample.docx"
    assert payload["results"][0]["status"] == "success"
    assert payload["results"][0]["article"]["title"] == "面向出版的智能结构化转换"
    assert payload["results"][1]["filename"] == "invalid.txt"
    assert payload["results"][1]["status"] == "failed"
    assert "仅支持 .docx 文件" in payload["results"][1]["error"]


def test_export_package_contains_required_files_and_media(tmp_path):
    media = tmp_path / "figure.png"
    media.write_bytes(b"png-content")
    from app.routers.convert import TEMP_ROOT

    package_media = TEMP_ROOT / "test-export" / "media" / "figure.png"
    package_media.parent.mkdir(parents=True, exist_ok=True)
    package_media.write_bytes(media.read_bytes())
    payload = {
        "filename": "sample.docx",
        "article": {
            "title": "打包测试",
            "abstract": "摘要",
            "keywords": ["ZIP", "JATS", "测试"],
            "sections": [{"title": "结果", "paragraphs": ["正文"]}],
            "figures": [{"id": "fig1", "caption": "图1", "path": str(package_media)}],
        },
        "xml": "<?xml version=\"1.0\"?><article/>",
        "media_paths": [str(package_media)],
        "validation": {"passed": True, "errors": [], "warnings": [], "xref_checks": []},
        "quality_report": {
            "total_score": 95,
            "grade": "A",
            "scores": {"metadata_score": 100},
            "issues": [],
            "summary": {"error_count": 0, "warning_count": 0, "suggestion_count": 0},
        },
    }

    response = client.post("/api/export-package", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/zip"
    with zipfile.ZipFile(BytesIO(response.content)) as archive:
        assert set(archive.namelist()) == {
            "article.xml",
            "article.json",
            "validation_report.md",
            "quality_report.json",
            "media/",
            "media/figure.png",
            "manifest.json",
        }
        article = json.loads(archive.read("article.json"))
        manifest = json.loads(archive.read("manifest.json"))
        assert article["title"] == "打包测试"
        assert manifest["source_filename"] == "sample.docx"
        assert manifest["quality_score"] == 95
        assert "media/figure.png" in manifest["files"]


def test_export_package_keeps_empty_media_directory():
    response = client.post(
        "/api/export-package",
        json={
            "article": {
                "title": "无媒体打包测试",
                "abstract": "摘要",
                "keywords": ["ZIP", "JATS", "测试"],
                "sections": [{"title": "结果", "paragraphs": ["正文"]}],
            },
            "xml": "<article/>",
            "media_paths": [],
            "validation": {"passed": True, "errors": [], "warnings": []},
        },
    )

    with zipfile.ZipFile(BytesIO(response.content)) as archive:
        assert "media/" in archive.namelist()


def test_export_package_rejects_media_outside_temp_root(tmp_path):
    outside = tmp_path / "secret.txt"
    outside.write_text("secret", encoding="utf-8")

    response = client.post(
        "/api/export-package",
        json={
            "article": {
                "title": "安全测试",
                "abstract": "摘要",
                "keywords": ["ZIP", "JATS", "测试"],
                "sections": [{"title": "结果", "paragraphs": ["正文"]}],
            },
            "xml": "<article/>",
            "media_paths": [str(outside)],
            "validation": {"passed": True, "errors": [], "warnings": []},
        },
    )

    assert response.status_code == 400
    assert "媒体路径不允许访问" in response.json()["detail"]
