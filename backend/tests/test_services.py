from io import BytesIO
import struct
import zlib

from docx import Document
from docx.shared import Inches
from lxml import etree

from app.services.docx_parser import DocxParser
from app.services.jats_generator import JatsGenerator
from app.services.validator import ArticleValidator


def make_png(color: tuple[int, int, int]) -> bytes:
    width, height = 8, 8
    row = bytes([0] + [channel for _ in range(width) for channel in (*color, 255)])
    raw = row * height

    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
        )

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )


def build_figure_docx(path, captions: list[str], image_count: int) -> None:
    document = Document()
    title = document.add_paragraph("图片提取测试")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("摘要：测试图片提取。")
    document.add_paragraph("关键词：图片；JATS；测试")
    document.add_paragraph("1 结果")
    for index in range(image_count):
        image_path = path.parent / f"source_{index + 1}.png"
        image_path.write_bytes(make_png((20 + index * 50, 100, 130)))
        document.add_picture(str(image_path), width=Inches(1))
    for caption in captions:
        document.add_paragraph(caption)
    document.save(path)


def build_sample_docx() -> bytes:
    document = Document()
    title = document.add_paragraph()
    title.alignment = 1
    run = title.add_run("面向出版的智能结构化转换")
    run.bold = True
    run.font.size = __import__("docx").shared.Pt(18)
    document.add_paragraph("张三，李四")
    document.add_paragraph("未来出版大学 智能出版研究院")
    document.add_paragraph("摘要：本文提出一种规则驱动的 Word 到 JATS 转换方法。")
    document.add_paragraph("关键词：JATS；Word；结构化出版")
    document.add_paragraph("1 引言")
    document.add_paragraph("学术出版需要可交换的结构化内容。")
    document.add_paragraph("（1）解析文档结构")
    document.add_paragraph("E = mc²")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 张三. 智能出版研究[J]. 2026.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_parser_extracts_core_article_structure(tmp_path):
    path = tmp_path / "sample.docx"
    path.write_bytes(build_sample_docx())

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["title"] == "面向出版的智能结构化转换"
    assert [author["name"] for author in article["authors"]] == ["张三", "李四"]
    assert article["affiliations"] == ["未来出版大学 智能出版研究院"]
    assert article["abstract"].startswith("本文提出")
    assert article["keywords"] == ["JATS", "Word", "结构化出版"]
    assert article["sections"][0]["title"] == "引言"
    assert article["lists"][0]["items"] == ["解析文档结构"]
    formula = article["formulas"][0]
    assert {
        key: formula[key]
        for key in (
            "id", "content", "omml", "mathml", "latex", "type",
            "conversion_status", "supported_features", "unsupported_features",
            "issues", "section_index",
        )
    } == {
        "id": "eq1",
        "content": "E = mc²",
        "omml": "",
        "mathml": "",
        "latex": "",
        "type": "plain_text",
        "conversion_status": "success",
        "supported_features": [],
        "unsupported_features": [],
        "issues": [],
        "section_index": 0,
    }
    assert formula["is_display"] is True
    assert formula["status"] == "need_review"
    assert formula["confidence"] >= 0.50
    assert article["references"][0]["id"] == "ref1"
    assert article["references"][0]["label"] == "[1]"
    assert article["references"][0]["raw"] == "张三. 智能出版研究[J]. 2026."
    assert article["references"][0]["article_title"] == "智能出版研究"


def test_generator_produces_parseable_jats():
    article = {
        "title": "测试文章",
        "authors": [{"name": "张三", "orcid": ""}],
        "affiliations": ["测试大学"],
        "abstract": "测试摘要",
        "keywords": ["JATS"],
        "sections": [{"title": "引言", "level": 1, "paragraphs": ["内容"]}],
        "figures": [],
        "lists": [],
        "formulas": [{"id": "eq1", "content": "E = mc²", "type": "plain_text", "section_index": 0}],
        "references": [{"raw": "[1] 测试参考文献"}],
    }

    xml = JatsGenerator().generate(article)
    result = ArticleValidator().validate(article, xml)

    assert "<article-title>测试文章</article-title>" in xml
    assert "<surname>张</surname>" in xml
    assert '<disp-formula id="eq1">' in xml
    assert "<![CDATA[E = mc²]]>" in xml
    assert '<xref ref-type="aff" rid="aff1"/>' in xml
    assert '<aff id="aff1">' in xml
    assert "<label>1</label> 测试大学" in xml
    assert result["passed"] is True


def test_generator_ignores_non_mathml_formula_fragment():
    article = {
        "title": "安全公式测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["JATS", "MathML", "安全"],
        "sections": [{"title": "方法", "paragraphs": ["正文"]}],
        "figures": [],
        "tables": [],
        "lists": [],
        "references": [],
        "formulas": [{
            "id": "eq1",
            "content": "x=1",
            "mathml": "<script xmlns='https://example.invalid'>alert(1)</script>",
            "section_index": 0,
        }],
    }

    xml = JatsGenerator().generate(article)

    assert "<script" not in xml
    assert "<![CDATA[x=1]]>" in xml


def test_parser_recognizes_reference_heading_and_common_labels(tmp_path):
    path = tmp_path / "references.docx"
    document = Document()
    title = document.add_paragraph("参考文献解析测试")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("摘要：测试参考文献解析。")
    document.add_paragraph("关键词：参考文献；JATS；测试")
    document.add_paragraph("1 引言")
    document.add_paragraph("正文内容。")
    document.add_paragraph(" References： ")
    document.add_paragraph("[1] Zhang S. Structured publishing[J]. 2025.")
    document.add_paragraph("2. Li S. JATS conversion[J]. 2026.")
    document.add_paragraph("（3）王五. 数字出版研究[J]. 2026.")
    document.add_paragraph("No label reference entry.")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [
        {key: reference[key] for key in ("id", "label", "raw")}
        for reference in article["references"]
    ] == [
        {
            "id": "ref1",
            "label": "[1]",
            "raw": "Zhang S. Structured publishing[J]. 2025.",
        },
        {
            "id": "ref2",
            "label": "2.",
            "raw": "Li S. JATS conversion[J]. 2026.",
        },
        {
            "id": "ref3",
            "label": "（3）",
            "raw": "王五. 数字出版研究[J]. 2026.",
        },
        {
            "id": "ref4",
            "label": "",
            "raw": "No label reference entry.",
        },
    ]


def test_generator_builds_labeled_reference_list():
    article = {
        "title": "参考文献测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["参考文献", "JATS", "测试"],
        "sections": [{"title": "引言", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "lists": [],
        "formulas": [],
        "references": [
            {"id": "ref1", "label": "[1]", "raw": "First citation."},
            {"id": "ref2", "label": "2.", "raw": "Second citation."},
        ],
    }

    xml = JatsGenerator().generate(article)

    assert "<ref-list>" in xml
    assert "<title>References</title>" in xml
    assert '<ref id="ref1">' in xml
    assert "<label>[1]</label>" in xml
    assert "<mixed-citation>First citation.</mixed-citation>" in xml


def test_generator_adds_affiliation_labels_when_text_has_no_label():
    article = {
        "title": "Affiliation labels",
        "authors": [{"name": "Alice Smith", "orcid": "", "affiliation_ids": ["aff1"]}],
        "affiliations": ["Faculty of Medicine, University of Antwerp, 2650 Antwerp"],
        "abstract": "Abstract",
        "keywords": ["JATS", "affiliation", "metadata"],
        "sections": [{"title": "Introduction", "level": 1, "paragraphs": ["Body"]}],
        "figures": [], "tables": [], "lists": [], "formulas": [], "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert '<aff id="aff1">' in xml
    assert "<label>1</label>" in xml
    assert "Faculty of Medicine, University of Antwerp" in xml


def test_generator_adds_labels_for_unnumbered_top_level_sections():
    article = {
        "title": "Section labels",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["JATS", "section", "labels"],
        "sections": [
            {"title": "Introduction", "level": 1, "paragraphs": ["Intro"]},
            {"title": "Methods", "level": 1, "paragraphs": ["Methods"]},
        ],
        "figures": [], "tables": [], "lists": [], "formulas": [], "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert "<label>1.</label>" in xml
    assert "<title>Introduction</title>" in xml
    assert "<label>2.</label>" in xml
    assert "<title>Methods</title>" in xml


def test_generator_splits_numbered_section_titles_into_label_and_title():
    article = {
        "title": "Numbered sections",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["JATS", "section", "labels"],
        "sections": [
            {"title": "1. Introduction", "level": 1, "paragraphs": ["Intro"]},
            {"title": "1.1 Study Design", "level": 2, "paragraphs": ["Design"]},
        ],
        "figures": [], "tables": [], "lists": [], "formulas": [], "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert "<label>1.</label>" in xml
    assert "<title>Introduction</title>" in xml
    assert "<label>1.1</label>" in xml
    assert "<title>Study Design</title>" in xml
    assert "<title>1. Introduction</title>" not in xml


def test_generator_builds_jats_publishing_metadata_and_affiliation_links():
    article = {
        "title": "JATS Publishing 测试",
        "doi": "10.1234/word2jats.2026.001",
        "article_type": "research-article",
        "lang": "zh",
        "journal_title": "智能出版研究",
        "journal_id": "W2J",
        "publisher_name": "未来出版学会",
        "subject": "数字出版",
        "pub_year": "2026",
        "pub_month": "06",
        "pub_day": "10",
        "authors": [
            {"name": "张三", "orcid": "0000-0001-2345-6789", "affiliation_ids": ["aff1"]},
            {"name": "李四", "orcid": "", "affiliation_ids": ["aff2"]},
        ],
        "affiliations": ["未来出版大学", "智能出版研究院"],
        "abstract": "摘要",
        "keywords": ["JATS", "Publishing", "测试"],
        "sections": [{"title": "引言", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert '<!DOCTYPE article PUBLIC "-//NLM//DTD JATS (Z39.96) Journal Publishing DTD with MathML3 v1.3 20210610//EN" "JATS-journalpublishing1-3-mathml3.dtd">' in xml
    assert 'article-type="research-article"' in xml
    assert 'dtd-version="1.3"' in xml
    assert 'xml:lang="zh"' in xml
    assert 'xmlns:xlink="http://www.w3.org/1999/xlink"' in xml
    assert "<journal-meta>" in xml
    assert '<journal-id journal-id-type="publisher-id">W2J</journal-id>' in xml
    assert "<journal-title>智能出版研究</journal-title>" in xml
    assert "<publisher-name>未来出版学会</publisher-name>" in xml
    assert '<article-id pub-id-type="doi">10.1234/word2jats.2026.001</article-id>' in xml
    assert "<subject>数字出版</subject>" in xml
    assert '<xref ref-type="aff" rid="aff1"/>' in xml
    assert '<xref ref-type="aff" rid="aff2"/>' in xml
    assert '<aff id="aff1">' in xml
    assert "<label>1</label> 未来出版大学" in xml
    assert '<aff id="aff2">' in xml
    assert "<label>2</label> 智能出版研究院" in xml
    assert '<pub-date publication-format="electronic">' in xml
    assert "<year>2026</year>" in xml


def test_generator_emits_issn_for_formal_schema_delivery():
    article = {
        "title": "ISSN test", "abstract": "Abstract", "keywords": ["a", "b", "c"],
        "sections": [{"title": "Intro", "paragraphs": ["Text"]}],
        "authors": [{"name": "Alice Smith", "orcid": "0000-0000-0000-0001"}],
        "affiliations": ["Publishing Lab"], "figures": [], "tables": [],
        "lists": [], "formulas": [], "references": [], "issn": "1234-5678",
    }

    xml = JatsGenerator().generate(article)

    assert '<issn pub-type="epub">1234-5678</issn>' in xml


def test_generator_emits_jats_13_xlink_graphics():
    article = {
        "title": "Figure test", "abstract": "Abstract", "keywords": ["a", "b", "c"],
        "sections": [{"title": "Intro", "paragraphs": ["Text"]}],
        "authors": [], "affiliations": [], "figures": [
            {"id": "fig1", "caption": "Fig. 1", "path": "media/figure.png", "section_index": 0}
        ],
        "tables": [], "lists": [], "formulas": [], "references": [], "issn": "1234-5678",
    }

    xml = JatsGenerator().generate(article)

    assert '<graphic xlink:href="media/figure.png"/>' in xml
    assert ' href="media/figure.png"' not in xml


def test_generator_omits_empty_jats_groups_for_schema_validity():
    article = {
        "title": "Minimal",
        "abstract": "Abstract",
        "keywords": [],
        "sections": [{"title": "Intro", "paragraphs": ["Text"]}],
        "authors": [],
        "affiliations": [],
        "figures": [],
        "tables": [],
        "lists": [],
        "formulas": [],
        "references": [],
        "journal_id": "W2J",
        "journal_title": "Word2JATS Demo Journal",
        "publisher_name": "Word2JATS Publishing Lab",
        "issn": "1234-5678",
    }

    xml = JatsGenerator().generate(article)

    assert "<contrib-group" not in xml
    assert "<kwd-group" not in xml
    assert "<ref-list" not in xml


def test_generator_assigns_stable_section_ids_for_flow_mapping():
    article = {
        "title": "Section ID test", "abstract": "Abstract", "keywords": ["a", "b", "c"],
        "sections": [{"title": "Intro", "paragraphs": ["Text"]}],
        "authors": [], "affiliations": [], "figures": [], "tables": [],
        "lists": [], "formulas": [], "references": [],
    }

    xml = JatsGenerator().generate(article)

    root = etree.fromstring(xml.encode("utf-8"))
    assert root.find("./body/sec").get("id") == "sec1"


def test_parser_recognizes_formula_symbols_keywords_and_equation_style(tmp_path):
    path = tmp_path / "formulas.docx"
    document = Document()
    title = document.add_paragraph("公式识别测试")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("摘要：测试基础公式识别。")
    document.add_paragraph("关键词：公式；JATS；测试")
    document.add_paragraph("1 方法")
    document.add_paragraph("x ≈ α + β")
    document.add_paragraph(r"\frac{a}{b} = sqrt(c)")
    styled = document.add_paragraph("customEquationContent")
    styled.style = document.styles.add_style("Equation Custom", 1)
    document.add_paragraph(
        "这是一段较长的普通正文，其中提到了 log 和 sin，但不应该作为独立公式识别，"
        "因为它超过了基础规则允许的公式段落长度。"
    )
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [formula["id"] for formula in article["formulas"]] == ["eq1", "eq2", "eq3"]
    assert [formula["content"] for formula in article["formulas"]] == [
        "x ≈ α + β",
        r"\frac{a}{b} = sqrt(c)",
        "customEquationContent",
    ]
    assert all(formula["type"] == "plain_text" for formula in article["formulas"])
    assert all(formula["section_index"] == 0 for formula in article["formulas"])


def test_generator_places_formula_in_section_and_uses_cdata():
    article = {
        "title": "公式测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["公式", "JATS", "测试"],
        "sections": [{"title": "方法", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "lists": [],
        "formulas": [
            {
                "id": "eq1",
                "content": r"\frac{a}{b} <= c & d",
                "type": "plain_text",
                "section_index": 0,
            }
        ],
        "references": [],
    }

    xml = JatsGenerator().generate(article)

    section_start = xml.index("<sec ")
    section_end = xml.index("</sec>")
    formula_index = xml.index('<disp-formula id="eq1">')
    assert section_start < formula_index < section_end
    assert r"<![CDATA[\frac{a}{b} <= c & d]]>" in xml


def test_generator_places_unlocated_formula_at_body_end():
    article = {
        "title": "公式测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["公式", "JATS", "测试"],
        "sections": [{"title": "方法", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "lists": [],
        "formulas": [
            {
                "id": "eq1",
                "content": "x = 1",
                "type": "plain_text",
                "section_index": -1,
            }
        ],
        "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert xml.index("</sec>") < xml.index('<disp-formula id="eq1">') < xml.index("</body>")


def test_parser_extracts_zip_media_and_binds_supported_captions(tmp_path):
    path = tmp_path / "figures.docx"
    build_figure_docx(path, ["图1-1 系统架构", "Fig. 2 Overview"], 2)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [figure["id"] for figure in article["figures"]] == ["fig1", "fig2"]
    assert [figure["caption"] for figure in article["figures"]] == [
        "图1-1 系统架构",
        "Fig. 2 Overview",
    ]
    assert article["figures"][0]["path"].replace("\\", "/").endswith("media/figure_1.png")
    assert (tmp_path / article["figures"][0]["path"]).exists()


def test_parser_keeps_extra_images_and_caption_only_figures(tmp_path):
    extra_image_docx = tmp_path / "extra-images.docx"
    build_figure_docx(extra_image_docx, ["图 1 第一张图片"], 2)
    image_article = DocxParser(extra_image_docx, tmp_path / "image-media").parse()

    assert len(image_article["figures"]) == 2
    assert image_article["figures"][1]["caption"] == ""
    assert image_article["figures"][1]["path"]

    extra_caption_docx = tmp_path / "extra-captions.docx"
    build_figure_docx(
        extra_caption_docx,
        ["Figure 1 Architecture", "图 2 实验流程"],
        1,
    )
    caption_article = DocxParser(extra_caption_docx, tmp_path / "caption-media").parse()

    assert len(caption_article["figures"]) == 2
    assert caption_article["figures"][1]["caption"] == "图 2 实验流程"
    assert caption_article["figures"][1]["path"] == ""


def test_generator_omits_graphic_for_caption_only_figure():
    article = {
        "title": "图片测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["图片", "JATS", "测试"],
        "sections": [{"title": "结果", "level": 1, "paragraphs": ["正文"]}],
        "figures": [
            {
                "id": "fig1",
                "caption": "Figure 1 Caption only",
                "path": "",
                "section_index": 0,
            }
        ],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert '<fig id="fig1">' in xml
    assert "<label>Figure 1</label>" in xml
    assert "<p>Caption only</p>" in xml
    assert "<graphic" not in xml


def test_parser_extracts_tables_and_binds_supported_captions(tmp_path):
    path = tmp_path / "tables.docx"
    document = Document()
    title = document.add_paragraph("表格解析测试")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("摘要：测试表格解析。")
    document.add_paragraph("关键词：表格；JATS；测试")
    document.add_paragraph("1 结果")
    table1 = document.add_table(rows=2, cols=3)
    for cell, value in zip(table1.rows[0].cells, ["指标", "方法A", "方法B"]):
        cell.text = value
    for cell, value in zip(table1.rows[1].cells, ["准确率", "90%", "95%"]):
        cell.text = value
    document.add_paragraph("表1 实验结果")
    table2 = document.add_table(rows=2, cols=2)
    for cell, value in zip(table2.rows[0].cells, ["Metric", "Value"]):
        cell.text = value
    for cell, value in zip(table2.rows[1].cells, ["Recall", "88%"]):
        cell.text = value
    document.add_paragraph("Table 2 Evaluation")
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert [
        {
            key: table[key]
            for key in ("id", "caption", "rows", "section_index")
        }
        for table in article["tables"]
    ] == [
        {
            "id": "tab1",
            "caption": "表1 实验结果",
            "rows": [["指标", "方法A", "方法B"], ["准确率", "90%", "95%"]],
            "section_index": 0,
        },
        {
            "id": "tab2",
            "caption": "Table 2 Evaluation",
            "rows": [["Metric", "Value"], ["Recall", "88%"]],
            "section_index": 0,
        },
    ]
    assert all(table["status"] == "ok" for table in article["tables"])
    assert all(table["confidence"] >= 0.80 for table in article["tables"])


def test_parser_extracts_table_from_document_without_non_empty_paragraphs(tmp_path):
    path = tmp_path / "table-only.docx"
    document = Document()
    document.add_table(rows=1, cols=2).rows[0].cells[0].text = "A"
    document.tables[0].rows[0].cells[1].text = "B"
    document.save(path)

    article = DocxParser(path, tmp_path / "media").parse()

    assert article["tables"][0]["rows"] == [["A", "B"]]


def test_parser_keeps_extra_tables_and_caption_only_tables(tmp_path):
    extra_table_path = tmp_path / "extra-table.docx"
    document = Document()
    document.add_paragraph("表格测试")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "A"
    document.add_paragraph("表 1 第一张表")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "B"
    document.save(extra_table_path)

    table_article = DocxParser(extra_table_path, tmp_path / "table-media").parse()
    assert len(table_article["tables"]) == 2
    assert table_article["tables"][1]["caption"] == ""
    assert table_article["tables"][1]["rows"] == [["B"]]

    extra_caption_path = tmp_path / "extra-caption.docx"
    document = Document()
    document.add_paragraph("表格测试")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "A"
    document.add_paragraph("表 1 第一张表")
    document.add_paragraph("Table 2 Caption only")
    document.save(extra_caption_path)

    caption_article = DocxParser(extra_caption_path, tmp_path / "caption-media").parse()
    assert len(caption_article["tables"]) == 2
    assert caption_article["tables"][1]["caption"] == "Table 2 Caption only"
    assert caption_article["tables"][1]["rows"] == []


def test_generator_builds_jats_table_wrap():
    article = {
        "title": "表格测试",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["表格", "JATS", "测试"],
        "sections": [{"title": "结果", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "tables": [
            {
                "id": "tab1",
                "caption": "表1 实验结果",
                "rows": [["指标", "方法A"], ["准确率", "95%"]],
                "section_index": 0,
            }
        ],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    xml = JatsGenerator().generate(article)

    assert '<table-wrap id="tab1">' in xml
    assert "<label>表1</label>" in xml
    assert "<p>实验结果</p>" in xml
    assert "<thead>" in xml
    assert "<th>指标</th>" in xml
    assert "<tbody>" in xml
    assert "<td>95%</td>" in xml


def test_generator_builds_table_wrap_foot_for_notes():
    article = {
        "title": "Table notes",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["table", "JATS", "notes"],
        "sections": [{"title": "Results", "level": 1, "paragraphs": ["Text"]}],
        "figures": [],
        "tables": [
            {
                "id": "tab1",
                "caption": "Table 1. Baseline metrics",
                "rows": [["Metric", "Value"], ["Accuracy", "95%"]],
                "notes": ["Note: Values are shown as percentages."],
                "section_index": 0,
            }
        ],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    root = etree.fromstring(JatsGenerator().generate(article).encode("utf-8"))
    table_wrap = root.xpath("//*[local-name()='table-wrap']")[0]

    assert table_wrap.xpath(
        "string(./*[local-name()='table-wrap-foot']/*[local-name()='fn']/*[local-name()='p'])"
    ) == "Note: Values are shown as percentages."
    assert [etree.QName(child).localname for child in table_wrap][-1] == "table-wrap-foot"


def test_generator_uses_tbody_for_single_row_table():
    article = {
        "title": "Single row table",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["table", "JATS", "DTD"],
        "sections": [{"title": "Results", "level": 1, "paragraphs": ["Text"]}],
        "figures": [],
        "tables": [{
            "id": "tab1",
            "caption": "Table 1. One row",
            "rows": [["Only cell"]],
            "section_index": 0,
        }],
        "lists": [],
        "formulas": [],
        "references": [],
        "journal_id": "W2J",
        "journal_title": "Word2JATS Test Journal",
        "publisher_name": "Word2JATS",
        "issn": "1234-5678",
    }

    xml = JatsGenerator().generate(article)
    root = etree.fromstring(xml.encode("utf-8"))

    assert not root.xpath("//table/thead")
    assert root.xpath("string(//table/tbody/tr/td)") == "Only cell"
    assert ArticleValidator().schema_validator.validate(xml)["jats_schema_valid"] is True


def test_generator_separates_float_labels_from_caption_text():
    article = {
        "title": "Float labels",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["float", "JATS", "labels"],
        "sections": [{"title": "Results", "level": 1, "paragraphs": ["Text"]}],
        "figures": [{
            "id": "fig1",
            "caption": "Fig. 1: Architecture",
            "path": "image.png",
            "section_index": 0,
        }],
        "tables": [{
            "id": "tab1",
            "caption": "Table 1. Results",
            "rows": [["Metric"], ["95%"]],
            "section_index": 0,
        }],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    root = etree.fromstring(JatsGenerator().generate(article).encode("utf-8"))

    assert root.xpath("string(//fig/label)") == "Fig. 1"
    assert root.xpath("string(//fig/caption/p)") == "Architecture"
    assert root.xpath("string(//table-wrap/label)") == "Table 1"
    assert root.xpath("string(//table-wrap/caption/p)") == "Results"


def test_generator_omits_empty_table_element_for_caption_only_table():
    article = {
        "title": "Table caption only",
        "authors": [],
        "affiliations": [],
        "abstract": "Abstract",
        "keywords": ["table", "JATS", "test"],
        "sections": [{"title": "Results", "level": 1, "paragraphs": ["Text"]}],
        "figures": [],
        "tables": [
            {
                "id": "tab1",
                "caption": "Table 1 Caption only",
                "rows": [],
                "section_index": 0,
            }
        ],
        "lists": [],
        "formulas": [],
        "references": [],
        "journal_id": "W2J",
        "journal_title": "Word2JATS Demo Journal",
        "publisher_name": "Word2JATS Publishing Lab",
        "issn": "1234-5678",
    }

    xml = JatsGenerator().generate(article)

    assert '<table-wrap id="tab1">' in xml
    assert "<label>Table 1</label>" in xml
    assert "<p>Caption only</p>" in xml
    assert "<table/>" not in xml


def test_validator_reports_required_content():
    article = {
        "title": "",
        "authors": [],
        "affiliations": [],
        "abstract": "",
        "keywords": [],
        "sections": [],
        "figures": [{"id": "fig1", "caption": "", "path": "image.png"}],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    result = ArticleValidator().validate(article, "<article />")

    assert result["passed"] is False
    assert len(result["errors"]) == 9
    assert "缺少 JATS journal-meta 节点。" in result["errors"]
    assert "缺少 JATS article-meta 节点。" in result["errors"]
    assert "缺少 JATS title-group 节点。" in result["errors"]
    assert "缺少 JATS body 节点。" in result["errors"]
    assert "缺少 JATS back 节点。" in result["errors"]
    assert "作者为空，建议补充作者信息。" in result["warnings"]
    assert "单位为空，建议补充作者单位。" in result["warnings"]
    assert "参考文献为空，建议补充参考文献。" in result["warnings"]
    assert "图片 fig1 缺少图题。" in result["warnings"]
    assert "关键词少于 3 个，建议补充关键词。" in result["warnings"]


def test_validator_reports_jats_quality_warnings():
    article = {
        "title": "测试文章",
        "authors": [{"name": "张三", "orcid": ""}],
        "affiliations": ["测试大学"],
        "abstract": "测试摘要",
        "keywords": ["JATS", "XML"],
        "sections": [{"title": "引言", "level": 1, "paragraphs": []}],
        "figures": [{"id": "fig1", "caption": "", "path": "image.png"}],
        "lists": [],
        "tables": [
            {"id": "tab1", "caption": "", "rows": [["A"]], "section_index": 0},
            {"id": "tab2", "caption": "表2 空表", "rows": [], "section_index": 0},
        ],
        "formulas": [{"id": "eq1", "content": "", "type": "plain_text", "section_index": 0}],
        "references": [],
    }
    xml = JatsGenerator().generate(article)

    result = ArticleValidator().validate(article, xml)

    assert result["passed"] is True
    assert "表格 tab1 缺少表题。" in result["warnings"]
    assert "表格 tab2 没有数据行。" in result["warnings"]
    assert "作者 张三 缺少 ORCID。" in result["warnings"]
    assert "章节“引言”没有正文段落。" in result["warnings"]
    assert "公式 eq1 内容为空。" in result["warnings"]
    assert "关键词少于 3 个，建议补充关键词。" in result["warnings"]


def test_validator_reports_invalid_xml_without_jats_node_duplicates():
    article = {
        "title": "测试文章",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["一", "二", "三"],
        "sections": [{"title": "引言", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    result = ArticleValidator().validate(article, "<article>")

    assert result["passed"] is False
    assert len([error for error in result["errors"] if "XML 无法解析" in error]) == 1
    assert not any("article-meta" in error for error in result["errors"])
    assert not any("body" in error for error in result["errors"])


def test_validator_reports_missing_jats_publishing_structure():
    article = {
        "title": "测试文章",
        "authors": [],
        "affiliations": [],
        "abstract": "摘要",
        "keywords": ["一", "二", "三"],
        "sections": [{"title": "引言", "level": 1, "paragraphs": ["正文"]}],
        "figures": [],
        "lists": [],
        "formulas": [],
        "references": [],
    }

    result = ArticleValidator().validate(
        article,
        "<article><front><article-meta/></front><body/><back/></article>",
    )

    assert result["passed"] is False
    assert result["errors"] == [
        "缺少 JATS journal-meta 节点。",
        "缺少 JATS title-group 节点。",
    ]
