import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxParser:
    SECTION_PATTERNS = (
        re.compile(r"^(\d+(?:\.\d+)*)[\s、.]+(.+)$"),
        re.compile(r"^([一二三四五六七八九十]+)、\s*(.+)$"),
        re.compile(r"^[（(]([一二三四五六七八九十]+)[）)]\s*(.+)$"),
    )
    AFFILIATION_WORDS = (
        "大学", "学院", "研究院", "实验室", "中心",
        "school", "university", "institute", "laboratory", "center",
    )
    ABSTRACT_RE = re.compile(r"^\s*(摘要|摘\s*要|abstract)\s*[:：]?\s*", re.I)
    KEYWORD_RE = re.compile(r"^\s*(关键词|关\s*键\s*词|key\s*words?|keywords)\s*[:：]?\s*", re.I)
    REFERENCE_RE = re.compile(r"^\s*(?:参考\s*文献|references?)\s*[:：]?\s*$", re.I)
    REFERENCE_LABEL_RE = re.compile(
        r"^\s*(?P<label>\[\s*\d+\s*\]|\(\s*\d+\s*\)|（\s*\d+\s*）|\d+\s*[.．、])"
        r"\s*(?P<raw>.*)$"
    )
    FIGURE_RE = re.compile(
        r"^\s*(?:图\s*\d+(?:\s*[-－—.]\s*\d+)*|fig(?:ure)?\.?\s*\d+(?:\s*[-.]\s*\d+)*)"
        r"(?:\s+|[:：])?.*$",
        re.I,
    )
    LIST_RE = re.compile(r"^\s*(?:[（(]\d+[）)]|\d+[）)]|[-•·])\s*")
    TABLE_RE = re.compile(
        r"^\s*(?:表\s*\d+(?:\s*[-－—.]\s*\d+)*|table\s*\d+(?:\s*[-.]\s*\d+)*)"
        r"(?:\s+|[:：])?.*$",
        re.I,
    )
    FORMULA_RE = re.compile(
        r"(=|≈|≤|≥|∑|∫|√|[αβγλμσ]|\\?frac|\\?sqrt|\blim\b|\blog\b|\bsin\b|\bcos\b)",
        re.I,
    )

    def __init__(self, docx_path: str | Path, media_dir: str | Path):
        self.docx_path = Path(docx_path)
        self.media_dir = Path(media_dir)

    def parse(self) -> dict[str, Any]:
        document = Document(self.docx_path)
        paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
        article = self._empty_article()
        if not paragraphs:
            article["figures"] = self._extract_images([])
            article["tables"] = self._extract_tables(document, [])
            return article

        title_index = self._find_title_index(paragraphs)
        article["title"] = paragraphs[title_index].text.strip()
        author_indexes, affiliation_indexes = self._extract_front_matter(
            paragraphs, title_index, article
        )
        skipped = {title_index, *author_indexes, *affiliation_indexes}
        current_section: dict[str, Any] | None = None
        current_section_index = -1
        abstract_parts: list[str] = []
        in_abstract = False
        in_references = False
        captions: list[tuple[str, int]] = []
        table_captions: list[tuple[str, int]] = []

        for index, paragraph in enumerate(paragraphs):
            if index in skipped:
                continue
            text = paragraph.text.strip()

            if self.REFERENCE_RE.match(text):
                in_references = True
                in_abstract = False
                continue
            if in_references:
                article["references"].append(
                    self._parse_reference(text, len(article["references"]) + 1)
                )
                continue

            keyword_match = self.KEYWORD_RE.match(text)
            if keyword_match:
                in_abstract = False
                keyword_text = self.KEYWORD_RE.sub("", text, count=1)
                article["keywords"] = self._split_values(keyword_text)
                continue

            abstract_match = self.ABSTRACT_RE.match(text)
            if abstract_match:
                in_abstract = True
                abstract_text = self.ABSTRACT_RE.sub("", text, count=1)
                if abstract_text:
                    abstract_parts.append(abstract_text)
                continue

            if self.FIGURE_RE.match(text):
                in_abstract = False
                captions.append((text, current_section_index))
                continue
            if self.TABLE_RE.match(text):
                in_abstract = False
                table_captions.append((text, current_section_index))
                continue

            section = self._parse_section_title(text)
            if section:
                in_abstract = False
                article["sections"].append(section)
                current_section = section
                current_section_index = len(article["sections"]) - 1
                continue
            if in_abstract:
                abstract_parts.append(text)
                continue

            if self._is_list(paragraph, text):
                item = self.LIST_RE.sub("", text, count=1).strip()
                article["lists"].append({
                    "id": f"list{len(article['lists']) + 1}",
                    "items": [item or text],
                    "section_index": current_section_index,
                })
                continue
            if self._is_formula(paragraph, text):
                article["formulas"].append({
                    "id": f"eq{len(article['formulas']) + 1}",
                    "content": text,
                    "type": "plain_text",
                    "section_index": current_section_index,
                })
                continue
            if current_section is not None:
                current_section["paragraphs"].append(text)

        article["abstract"] = "\n".join(abstract_parts)
        article["figures"] = self._extract_images(captions)
        article["tables"] = self._extract_tables(document, table_captions)
        return article

    @staticmethod
    def _empty_article() -> dict[str, Any]:
        return {
            "title": "", "authors": [], "affiliations": [], "abstract": "",
            "keywords": [], "sections": [], "figures": [], "lists": [],
            "tables": [], "formulas": [], "references": [],
        }

    def _find_title_index(self, paragraphs: list[Any]) -> int:
        best_index = 0
        best_score = -1.0
        for index, paragraph in enumerate(paragraphs[:10]):
            text = paragraph.text.strip()
            score = 0.0
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                score += 4
            if any(run.bold for run in paragraph.runs):
                score += 3
            sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
            if sizes:
                score += min(max(sizes) / 4, 5)
            if 4 <= len(text) <= 45:
                score += 2
            if self.ABSTRACT_RE.match(text) or self.KEYWORD_RE.match(text):
                score -= 8
            if score > best_score:
                best_index, best_score = index, score
        return best_index

    def _extract_front_matter(
        self, paragraphs: list[Any], title_index: int, article: dict[str, Any]
    ) -> tuple[set[int], set[int]]:
        author_indexes: set[int] = set()
        affiliation_indexes: set[int] = set()
        window = range(title_index + 1, min(len(paragraphs), title_index + 7))
        for index in window:
            text = paragraphs[index].text.strip()
            lower = text.lower()
            if (
                self.ABSTRACT_RE.match(text)
                or self.KEYWORD_RE.match(text)
                or self.REFERENCE_RE.match(text)
                or self._parse_section_title(text)
            ):
                break
            if any(word in lower for word in self.AFFILIATION_WORDS):
                article["affiliations"].append(text)
                affiliation_indexes.add(index)
                continue
            if (
                not author_indexes
                and len(text) <= 80
                and not self.ABSTRACT_RE.match(text)
                and not self.KEYWORD_RE.match(text)
                and not self.FIGURE_RE.match(text)
                and not self.TABLE_RE.match(text)
                and not self._is_formula(paragraphs[index], text)
                and not self._parse_section_title(text)
                and re.search(r"[，,；;\s、]", text)
            ):
                names = [name for name in re.split(r"[，,；;\s、]+", text) if name]
                article["authors"] = [{"name": name, "orcid": ""} for name in names]
                author_indexes.add(index)
        return author_indexes, affiliation_indexes

    def _parse_reference(self, text: str, index: int) -> dict[str, str]:
        match = self.REFERENCE_LABEL_RE.match(text)
        label = match.group("label").strip() if match else ""
        raw = match.group("raw").strip() if match else text.strip()
        return {"id": f"ref{index}", "label": label, "raw": raw}

    def _parse_section_title(self, text: str) -> dict[str, Any] | None:
        for pattern in self.SECTION_PATTERNS:
            match = pattern.match(text)
            if not match:
                continue
            marker, title = match.groups()
            level = marker.count(".") + 1 if marker[0].isdigit() else (
                2 if text.startswith(("（", "(")) else 1
            )
            return {"title": title.strip(), "level": level, "paragraphs": []}
        return None

    def _is_list(self, paragraph: Any, text: str) -> bool:
        if self.LIST_RE.match(text):
            return True
        properties = paragraph._p.pPr
        return bool(properties is not None and properties.numPr is not None)

    def _is_formula(self, paragraph: Any, text: str) -> bool:
        style_name = paragraph.style.name if paragraph.style else ""
        styled_as_equation = "equation" in style_name.lower() or "公式" in style_name
        return styled_as_equation or (len(text) <= 60 and bool(self.FORMULA_RE.search(text)))

    @staticmethod
    def _split_values(text: str) -> list[str]:
        return [value.strip() for value in re.split(r"[；;，,、]+", text) if value.strip()]

    def _extract_images(self, captions: list[tuple[str, int]]) -> list[dict[str, Any]]:
        self.media_dir.mkdir(parents=True, exist_ok=True)
        image_paths: list[str] = []
        with zipfile.ZipFile(self.docx_path) as archive:
            media_names = sorted(
                (
                    name
                    for name in archive.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                ),
                key=self._natural_sort_key,
            )
            for index, media_name in enumerate(media_names, start=1):
                suffix = Path(media_name).suffix.lower() or ".bin"
                filename = f"figure_{index}{suffix}"
                path = self.media_dir / filename
                path.write_bytes(archive.read(media_name))
                image_paths.append(self._relative_path(path))

        figures = []
        for index in range(max(len(image_paths), len(captions))):
            caption, section_index = captions[index] if index < len(captions) else ("", -1)
            image_path = image_paths[index] if index < len(image_paths) else ""
            figures.append({
                "id": f"fig{index + 1}",
                "caption": caption,
                "path": image_path,
                "section_index": section_index,
            })
        return figures

    @staticmethod
    def _extract_tables(
        document: Any, captions: list[tuple[str, int]]
    ) -> list[dict[str, Any]]:
        table_rows = [
            [[cell.text.strip() for cell in row.cells] for row in table.rows]
            for table in document.tables
        ]
        tables = []
        for index in range(max(len(table_rows), len(captions))):
            caption, section_index = captions[index] if index < len(captions) else ("", -1)
            rows = table_rows[index] if index < len(table_rows) else []
            tables.append({
                "id": f"tab{index + 1}",
                "caption": caption,
                "rows": rows,
                "section_index": section_index,
            })
        return tables

    @staticmethod
    def _natural_sort_key(value: str) -> list[Any]:
        return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", value)]

    def _relative_path(self, path: Path) -> str:
        try:
            return path.resolve().relative_to(Path.cwd().resolve()).as_posix()
        except ValueError:
            return (Path(self.media_dir.name) / path.name).as_posix()
